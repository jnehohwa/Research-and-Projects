"""Apply targeted review feedback to Joshua's formatted ITEPA report."""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Emu


ROOT = Path(__file__).resolve().parents[1]
SOURCE = Path("/Users/joshuanehohwa/Downloads/Joshua_Nehohwa_ITEPA333_formatted.docx")
OUTPUT = ROOT / "output" / "docx" / "Joshua_Nehohwa_ITEPA333_formatted_revised.docx"
SCREENSHOTS = ROOT / "evidence" / "screenshots"
EVIDENCE_WIDTH = Emu(6_299_200)


def find_paragraph(document: Document, exact_text: str) -> Paragraph:
    return next(paragraph for paragraph in document.paragraphs if paragraph.text == exact_text)


def previous_paragraph(document: Document, target: Paragraph) -> Paragraph:
    paragraphs = document.paragraphs
    index = next(index for index, paragraph in enumerate(paragraphs) if paragraph._p is target._p)
    return paragraphs[index - 1]


def copy_run_format(source_run, destination_run) -> None:
    if source_run._r.rPr is not None:
        destination_run._r.insert(0, deepcopy(source_run._r.rPr))


def new_paragraph_after(anchor: Paragraph, template: Paragraph) -> Paragraph:
    element = OxmlElement("w:p")
    if template._p.pPr is not None:
        element.append(deepcopy(template._p.pPr))
    anchor._p.addnext(element)
    return Paragraph(element, anchor._parent)


def insert_text_after(anchor: Paragraph, text: str, template: Paragraph) -> Paragraph:
    paragraph = new_paragraph_after(anchor, template)
    run = paragraph.add_run(text)
    copy_run_format(template.runs[-1], run)
    return paragraph


def insert_bullet_before(anchor: Paragraph, text: str, template: Paragraph) -> Paragraph:
    paragraph_element = deepcopy(template._p)
    anchor._p.addprevious(paragraph_element)
    paragraph = Paragraph(paragraph_element, anchor._parent)
    paragraph.runs[0].text = "-"
    paragraph.runs[1].text = text
    for run in list(paragraph.runs[2:]):
        paragraph._p.remove(run._r)
    return paragraph


def insert_bullet_after(anchor: Paragraph, text: str, template: Paragraph) -> Paragraph:
    paragraph_element = deepcopy(template._p)
    anchor._p.addnext(paragraph_element)
    paragraph = Paragraph(paragraph_element, anchor._parent)
    paragraph.runs[0].text = "-"
    paragraph.runs[1].text = text
    for run in list(paragraph.runs[2:]):
        paragraph._p.remove(run._r)
    return paragraph


def insert_image_after(
    anchor: Paragraph,
    image_path: Path,
    image_template: Paragraph,
    width: Emu = EVIDENCE_WIDTH,
) -> Paragraph:
    paragraph = new_paragraph_after(anchor, image_template)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=width)
    return paragraph


def insert_caption_after(
    anchor: Paragraph, text: str, caption_template: Paragraph
) -> Paragraph:
    paragraph = new_paragraph_after(anchor, caption_template)
    run = paragraph.add_run(text)
    copy_run_format(caption_template.runs[0], run)
    paragraph.alignment = caption_template.alignment
    return paragraph


def main() -> None:
    document = Document(SOURCE)

    abstraction = find_paragraph(
        document,
        "-Abstraction: assessment calculations depend on the AssessmentStrategy interface rather than a specific algorithm.",
    )
    combined = find_paragraph(
        document,
        "-Polymorphism: each strategy provides the same calculate interface but returns its algorithm-specific result.-Single responsibility: registration workflow, monitoring and entity state are held in separate modules.",
    )
    combined.runs[0].text = "-"
    combined.runs[1].text = (
        "Polymorphism: each strategy provides the same calculate interface but returns "
        "its algorithm-specific result."
    )
    for run in list(combined.runs[2:]):
        combined._p.remove(run._r)
    insert_bullet_before(
        combined,
        "Inheritance: WeightedAverageStrategy, BestAttemptStrategy and PassFailStrategy "
        "inherit from the AssessmentStrategy abstract base class, reuse its score "
        "validation and provide their own calculate() implementations.",
        abstraction,
    )
    insert_bullet_after(
        combined,
        "Single responsibility: registration workflow, monitoring and entity state are "
        "held in separate modules.",
        abstraction,
    )

    pattern_explanation = find_paragraph(
        document,
        "The Singleton lock protects first-time creation if multiple threads request configuration simultaneously. The Factory rejects unsupported categories, while the Strategy implementations validate score and weight boundaries. The demonstration proves that both configuration requests return the same object identity and that each algorithm produces a distinct, correct result.",
    )
    insert_text_after(
        pattern_explanation,
        "Figure 1 maps each visible output to a pattern: 'Singleton: same instance = True' "
        "confirms one configuration instance; the four 'Factory' lines show category-based "
        "priority creation; and the weighted-average, best-attempt and pass/fail lines show "
        "the three interchangeable Strategy implementations.",
        pattern_explanation,
    )

    registration_evidence = find_paragraph(
        document,
        "The demonstration processes 16 requests, including a deliberate duplicate, against a course with capacity 10. Because request completion order is deliberately nondeterministic, the duplicate request may win the race and the earlier request may be reported as duplicate; the invariant remains that only one L001-PY701 record exists.",
    )
    insert_text_after(
        registration_evidence,
        "Figure 2 provides the complete execution evidence for this engine: all 16 request "
        "outcomes, the processing summary and the final integrity check are shown together.",
        registration_evidence,
    )

    concurrency_final = find_paragraph(
        document,
        "this sequence atomically prevents both check-then-act races and lost capacity updates.",
    )
    insert_text_after(
        concurrency_final,
        "Lock-order verification: _process_atomically() releases RegistrationService._lock "
        "before measure_transaction() reaches its final event-recording step. Bugzot therefore "
        "takes its own lock only after the registration lock has been released, so these two "
        "locks are not held at the same time in this workflow.",
        registration_evidence,
    )

    monitoring_export = find_paragraph(
        document,
        "Events can be exported to CSV for investigation and to JSON for management reporting. Bugzot itself is independent of console presentation, which keeps monitoring reusable.",
    )
    monitoring_heading = find_paragraph(document, "3.2 Application performance monitoring")
    aggregate_caption = find_paragraph(
        document, "Figure 3. Bugzot performance report generated from the registration run."
    )
    monitoring_image_template = previous_paragraph(document, aggregate_caption)
    raw_intro = insert_text_after(
        monitoring_export,
        "The evidence extract below shows confirmed, capacity-rejected and duplicate-rejected "
        "events with populated severity, context, correlation, timing and UTC timestamp fields.",
        monitoring_export,
    )
    raw_image = insert_image_after(
        raw_intro,
        SCREENSHOTS / "bugzot_raw_events.png",
        monitoring_image_template,
    )
    insert_caption_after(
        raw_image,
        "Figure 3A. Raw Bugzot events showing the diagnostic fields captured for individual transactions.",
        aggregate_caption,
    )
    aggregate_caption.runs[0].text = (
        "Figure 3B. Bugzot performance report generated from the registration run."
    )

    testing_intro = find_paragraph(
        document,
        "Pytest was selected because it provides concise tests, parameterisation, clear failure output and a mature coverage ecosystem. The 32-test suite covers normal behaviour, boundaries, invalid inputs, every required pattern, business rules, concurrent invariants, structured monitoring and the benchmark harness.",
    )
    coverage_caption = find_paragraph(
        document, "Figure 10. Successful automated test and coverage run."
    )
    test_image_template = previous_paragraph(document, coverage_caption)
    test_intro = insert_text_after(
        testing_intro,
        "The representative parameterised test below supplies four invalid email inputs and "
        "asserts that each one raises ValidationError. This demonstrates reusable boundary and "
        "invalid-input testing rather than only reporting the final pass count.",
        testing_intro,
    )
    test_image = insert_image_after(
        test_intro,
        SCREENSHOTS / "representative_test_case.png",
        test_image_template,
    )
    insert_caption_after(
        test_image,
        "Figure 10A. Representative parameterised pytest case for invalid learner email boundaries.",
        coverage_caption,
    )
    coverage_caption.runs[0].text = (
        "Figure 10B. Successful automated test and coverage run."
    )

    # Keep the next headings attached to their own sections after the inserted evidence.
    monitoring_heading.paragraph_format.keep_with_next = True

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
