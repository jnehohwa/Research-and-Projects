from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SOURCE = Path('/Users/joshuanehohwa/Downloads/ITOPA3-33 - Assignment - Mowbray - EDUV4948467.docx')
OUTPUT = Path('/Users/joshuanehohwa/Downloads/ITOPA3-33 - Assignment - Mowbray - EDUV4948467 - Revised.docx')


def paragraph_starting(document, prefix):
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text.startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f'Expected one paragraph starting with {prefix!r}, found {len(matches)}')
    return matches[0]


def paragraph_exact(document, text):
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text == text]
    if len(matches) != 1:
        raise RuntimeError(f'Expected one paragraph equal to {text!r}, found {len(matches)}')
    return matches[0]


def replace_runs(paragraph, text):
    run_properties = None
    if paragraph.runs and paragraph.runs[0]._r.rPr is not None:
        run_properties = deepcopy(paragraph.runs[0]._r.rPr)
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    new_run = paragraph.add_run(text)
    if run_properties is not None:
        new_run._r.insert(0, run_properties)


def append_like_last_run(paragraph, text):
    run_properties = None
    if paragraph.runs and paragraph.runs[-1]._r.rPr is not None:
        run_properties = deepcopy(paragraph.runs[-1]._r.rPr)
    new_run = paragraph.add_run(text)
    if run_properties is not None:
        new_run._r.insert(0, run_properties)


def replace_labelled_bullet(paragraph, label, body):
    label_properties = deepcopy(paragraph.runs[0]._r.rPr) if paragraph.runs and paragraph.runs[0]._r.rPr is not None else None
    body_properties = deepcopy(paragraph.runs[-1]._r.rPr) if paragraph.runs and paragraph.runs[-1]._r.rPr is not None else None
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    label_run = paragraph.add_run(label)
    if label_properties is not None:
        label_run._r.insert(0, label_properties)
    body_run = paragraph.add_run(body)
    if body_properties is not None:
        body_run._r.insert(0, body_properties)


def replace_text_in_run(paragraph, old, new):
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return
    raise RuntimeError(f'Could not find {old!r} within one run of paragraph {paragraph.text!r}')


def insert_reference_before(target, text):
    new_paragraph = target.insert_paragraph_before()
    if target._p.pPr is not None:
        new_paragraph._p.insert(0, deepcopy(target._p.pPr))
    run = new_paragraph.add_run(text)
    if target.runs and target.runs[0]._r.rPr is not None:
        run._r.insert(0, deepcopy(target.runs[0]._r.rPr))
    return new_paragraph


document = Document(SOURCE)

# 1.1: explicitly connect the scenario to the four calls in sequence.
intro = paragraph_starting(document, 'When a customer requests a bank statement')
append_like_last_run(
    intro,
    " In sequence, the application would open the customer's statement file, seek to the relevant period, read the data and close the file, with each step mediated by the corresponding system call above.",
)

# 1.3.2: add the complete system-call execution mechanism.
justification = paragraph_starting(document, 'System calls are the controlled gateway')
replace_runs(
    justification,
    'System calls are the controlled gateway between an application running in user mode and privileged operating-system services. '
    'A call triggers a trap or software interrupt, switching the CPU mode bit from user mode to kernel mode. The interrupt vector and system-call number route the request through the system-call table to the correct kernel handler, which validates and executes it before returning a result or error code and restoring user mode. '
    'This matches the course distinction between user mode and kernel mode (Eduvos, 2026a, slide 41) and protects the application, hardware and other processes. '
    'In a regulated bank, confidentiality, correctness, recovery and audit trails outweigh minor call overhead. If profiling reveals an I/O bottleneck, the bank should use supported options such as asynchronous I/O, batching, caching, faster storage or an approved driver update rather than bypassing the OS.',
)

# 2.2: distinguish economy from resource sharing and tie each benefit to ShopFast.
replace_labelled_bullet(
    paragraph_starting(document, 'Responsiveness: one blocked thread'),
    'Economy:',
    ' creating and switching threads is cheaper than creating separate processes because less operating-system state and memory are required.',
)
replace_labelled_bullet(
    paragraph_exact(document, 'Economy and resource sharing: threads share code, process data and resources, and are cheaper to manage than separate processes.'),
    'Resource sharing:',
    ' threads in the same process share the address space, code, data and open file descriptors, making cooperation faster and simpler.',
)
replace_labelled_bullet(
    paragraph_starting(document, 'Scalability: runnable threads'),
    'Scalability:',
    ' runnable threads can execute in parallel across the four CPU cores (Eduvos, 2026b, slide 46).',
)
replace_labelled_bullet(
    paragraph_starting(document, 'Responsiveness: while one customer waits'),
    'Economy:',
    ' during Black Friday, reusing lightweight request threads reduces creation and context-switching overhead, leaving more CPU time and memory for customer requests.',
)
replace_labelled_bullet(
    paragraph_exact(document, "Economy and resource sharing: request threads can reuse ShopFast's configuration, connection pools and product cache, leaving more memory and processing capacity for customer work."),
    'Resource sharing:',
    " ShopFast's request threads can use the same configuration, product cache, connection pools and open resources without duplicating them for every customer.",
)
replace_labelled_bullet(
    paragraph_starting(document, 'Scalability: product browsing'),
    'Scalability:',
    ' product browsing, cart updates and payment requests can run on different cores at the same time, increasing peak throughput provided the database does not become the next bottleneck.',
)

# Q3: support deadlock, circular-wait prevention and Banker's algorithm with the prescribed text.
q31 = paragraph_starting(document, 'The graph contains two cycles:')
replace_text_in_run(q31, '(Eduvos, 2026c, slide 66)', '(Eduvos, 2026c, slide 66; Silberschatz, Galvin and Gagne, 2019)')
q32 = paragraph_starting(document, 'This rule removes the circular-wait condition')
replace_text_in_run(q32, '(Eduvos, 2026c, slides 65 and 70)', '(Eduvos, 2026c, slides 65 and 70; Silberschatz, Galvin and Gagne, 2019)')
q33 = paragraph_starting(document, 'Deadlock prevention is more suitable')
replace_text_in_run(q33, '(Eduvos, 2026c).', '(Eduvos, 2026c; Silberschatz, Galvin and Gagne, 2019).')

# 4.5b: lead with a clear position, then acknowledge the limited opposing case.
replace_runs(
    paragraph_starting(document, 'There is no universally better design'),
    'Application-managed typing is better for a general-purpose system, provided the operating system retains a small set of structural file types. '
    'Applications can support new formats through filename extensions, registered associations, MIME information or content signatures without requiring kernel changes, which improves flexibility and portability. '
    'The operating system should still distinguish directories, regular files, symbolic links and devices because those types affect protection and device control. '
    'Strong OS-managed typing is preferable in tightly controlled or safety-critical environments, but for an ordinary desktop or server system, application-managed detailed formats combined with OS-managed structural types provide the best balance of flexibility and safety.',
)

# Add the prescribed textbook to the reference list in alphabetical order.
open_group = paragraph_starting(document, 'The Open Group.')
insert_reference_before(
    open_group,
    "Silberschatz, A., Galvin, P.B. and Gagne, G. 2019. Silberschatz's Operating System Concepts, Global Edition. 10th ed. Wiley. ISBN 9781119454083. Available at: https://bcs.wiley.com/he-bcs/Books?action=index&bcsId=11582&itemId=1119454085 (Accessed 28 August 2026).",
)

# Refresh the cached TOC page results for the final layout. The cover sheet is
# unnumbered, so displayed document page numbers are one less than PDF pages.
toc_pages = {
    'Question 1 - Operating-system system calls': '3',
    '1.1 Basic file-management system calls and their importance': '3',
    '1.2 Consequences of direct hard-drive and printer access': '4',
    '1.3.1 CIO decision': '4',
    '1.3.2 Justification': '4',
    'Question 2 - Multithreaded ShopFast server': '5',
    '2.1a Why the single-threaded server performed poorly': '5',
    '2.1b Blocking and its effect on clients': '5',
    '2.2a Three benefits of multithreading': '5',
    '2.2b Direct application to ShopFast': '5',
    '2.3a Task parallelism compared with data parallelism': '6',
    '2.3b Classification of the ShopFast server': '6',
    '2.4a Thread pool': '6',
    '2.4b Why pooling is preferred': '6',
    '2.4c Role of the task queue': '6',
    'Question 3 - Deadlock in the manufacturing system': '8',
    '3.1 Resource Allocation Graph and deadlock result': '8',
    '3.2 Safe request order': '9',
    '3.3 Prevention or avoidance': '9',
    'Question 4 - File management and caching': '10',
    '4.1 Three advantages of directories': '10',
    '4.2 Full path and Command Prompt line': '10',
    '4.3 HR and Public contents after the actions': '10',
    '4.4 Delete-on-logout compared with keep-until-deleted': '11',
    '4.5a Why systems treat file types differently': '11',
    '4.5b Which system is better?': '12',
    '4.6a How caches improve performance': '12',
    '4.6b Why systems do not use unlimited or larger caches': '12',
    'Reference list': '13',
}
for toc_paragraph in document._element.body.xpath('.//w:sdtContent//w:p'):
    text_nodes = toc_paragraph.xpath('.//w:t')
    if len(text_nodes) >= 2:
        entry_text = text_nodes[0].text or ''
        if entry_text in toc_pages:
            text_nodes[-1].text = toc_pages[entry_text]

# Ask Word to refresh the table of contents and other fields when the file opens.
settings = document.settings._element
existing = settings.find(qn('w:updateFields'))
if existing is None:
    existing = OxmlElement('w:updateFields')
    settings.append(existing)
existing.set(qn('w:val'), 'true')

document.save(OUTPUT)
print(OUTPUT)
