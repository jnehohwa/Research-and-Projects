from __future__ import annotations

import importlib.util
import json
import re
from html import unescape
from html.parser import HTMLParser
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/joshuanehohwa/Documents/Research and Projects")
TMP = ROOT / "tmp/itopa_docx"
SOURCE_BUILDER = ROOT / "tmp/itopa_final/build_final_pdf.py"
OUTPUT = ROOT / "ITOPA Practical/ITOPA3-33 - Assignment - Mowbray - EDUV4948467.docx"
COVER_IMAGE = TMP / "cover-page.png"
GRAPH_IMAGE = TMP / "resource-allocation-graph.png"
TOC_JSON = TMP / "toc-pages.json"

BLUE = "1F4D78"
TABLE_BLUE = "DCE6F1"
GRID = "B8C4D4"
INK = RGBColor(24, 34, 48)


class InlineHTMLParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.bold = False
        self.italic = False
        self.parts: list[tuple[str, bool, bool]] = []

    def handle_starttag(self, tag, attrs):
        if tag == "b":
            self.bold = True
        elif tag == "i":
            self.italic = True
        elif tag == "br":
            self.parts.append(("\n", self.bold, self.italic))

    def handle_startendtag(self, tag, attrs):
        if tag == "br":
            self.parts.append(("\n", self.bold, self.italic))

    def handle_endtag(self, tag):
        if tag == "b":
            self.bold = False
        elif tag == "i":
            self.italic = False

    def handle_data(self, data):
        if data:
            self.parts.append((unescape(data), self.bold, self.italic))


def set_font(run, name="Carlito", size=12, bold=None, italic=None, color=INK):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=GRID, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_inches):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [round(w * 1440) for w in widths_inches]
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def set_bottom_border(paragraph, color=BLUE, size=10, space=4):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)


def add_page_field(paragraph):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "2"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, instr, separate, text, end])
    set_font(run, "Arial", 9, color=RGBColor(0, 0, 0))


def add_html(paragraph, html_text, base_size=12, font="Carlito", color=INK):
    parser = InlineHTMLParser()
    parser.feed(html_text)
    for text, bold, italic in parser.parts:
        if text == "\n":
            paragraph.add_run().add_break()
            continue
        run = paragraph.add_run(text)
        set_font(run, font, base_size, bold=bold, italic=italic, color=color)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Carlito"
    normal.font.size = Pt(12)
    normal.font.color.rgb = INK
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Carlito")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Carlito")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.first_line_indent = Inches(0.24)

    h1 = styles["Heading 1"]
    h1.font.name = "Carlito"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(31, 77, 120)
    h1._element.rPr.rFonts.set(qn("w:ascii"), "Carlito")
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), "Carlito")
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = "Carlito"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(31, 77, 120)
    h2._element.rPr.rFonts.set(qn("w:ascii"), "Carlito")
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), "Carlito")
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h2.paragraph_format.space_before = Pt(11)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Carlito"
        style.font.size = Pt(12)
        style.font.color.rgb = INK
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.space_after = Pt(4)


def load_pdf_story():
    spec = importlib.util.spec_from_file_location("itopa_pdf_builder", SOURCE_BUILDER)
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module, module.build_story()


def plain_from_reportlab_paragraph(cell):
    parser = InlineHTMLParser()
    parser.feed(getattr(cell, "text", ""))
    return "".join(text for text, _, _ in parser.parts).strip()


def add_reportlab_table(doc, flowable):
    rows = flowable._cellvalues
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    raw_widths = [float(w or 1) for w in flowable._colWidths]
    usable = 7.18
    total = sum(raw_widths)
    widths = [usable * w / total for w in raw_widths]
    set_table_geometry(table, widths)
    set_table_borders(table)
    for r_idx, row in enumerate(rows):
        for c_idx, source in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell, 90, 95, 90, 95)
            if r_idx == 0:
                set_cell_shading(cell, TABLE_BLUE)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            value = plain_from_reportlab_paragraph(source) if hasattr(source, "text") else str(source)
            run = p.add_run(value)
            set_font(run, "Carlito", 9.2, bold=(r_idx == 0), color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_bullet_list(doc, flowable):
    for item in flowable._flowables:
        for child in item._flowables:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.first_line_indent = None
            add_html(p, child.text, 12)


def add_body_paragraph(doc, text, style_name):
    if style_name == "Question":
        p = doc.add_paragraph(style="Heading 1")
        if text.startswith(("Question 2", "Question 3", "Question 4")):
            p.paragraph_format.page_break_before = True
        add_html(p, text, 16, "Carlito", RGBColor(31, 77, 120))
    elif style_name == "SubQuestion":
        p = doc.add_paragraph(style="Heading 2")
        add_html(p, text, 13, "Carlito", RGBColor(31, 77, 120))
    elif style_name == "AssignmentTitle":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        add_html(p, text, 18, "Carlito", INK)
        for run in p.runs:
            run.bold = True
        set_bottom_border(p)
    elif style_name == "CodeA":
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.right_indent = Inches(0.18)
        p.paragraph_format.space_after = Pt(6)
        add_html(p, text, 10.5, "Courier New", INK)
        p_pr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F3F5F7")
        p_pr.append(shd)
    elif style_name == "Reference":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(3)
        add_html(p, text, 9.2)
    elif style_name == "Callout":
        text = text.replace("formatted this PDF.", "formatted the PDF and this Word document.")
        table = doc.add_table(rows=1, cols=1)
        set_table_geometry(table, [7.18])
        set_table_borders(table, BLUE, 8)
        cell = table.cell(0, 0)
        set_cell_shading(cell, "E6EEF7")
        set_cell_margins(cell, 120, 130, 120, 130)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.line_spacing = 1.15
        add_html(p, text, 10.5)
    else:
        p = doc.add_paragraph()
        add_html(p, text, 12)
    return p


def add_graph(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run()
    shape = run.add_picture(str(GRAPH_IMAGE), width=Inches(7.05))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", "Resource Allocation Graph showing the deadlocked manufacturing processes and resources")


def process_flowable(doc, module, flowable):
    name = type(flowable).__name__
    if name == "Paragraph":
        add_body_paragraph(doc, flowable.text, flowable.style.name)
    elif name == "Table":
        add_reportlab_table(doc, flowable)
    elif name == "ListFlowable":
        add_bullet_list(doc, flowable)
    elif name == "KeepTogether":
        for child in flowable._content:
            process_flowable(doc, module, child)
    elif name == "ResourceAllocationGraph":
        add_graph(doc)
    elif name == "PageBreak":
        doc.add_page_break()
    elif name == "Spacer":
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
    elif name == "HRFlowable":
        return


def add_toc(doc, pages):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    set_bottom_border(p, size=8)
    run = p.add_run("Table of Contents")
    set_font(run, "Carlito", 16, bold=True, color=RGBColor(31, 77, 120))
    entries = [
        ("Question 1 - Operating-system system calls (20 marks)", pages.get("Question 1", 3)),
        ("Question 2 - Multithreaded ShopFast server (25 marks)", pages.get("Question 2", 5)),
        ("Question 3 - Deadlock in the manufacturing system (20 marks)", pages.get("Question 3", 7)),
        ("Question 4 - File management and caching (35 marks)", pages.get("Question 4", 9)),
        ("Reference list", pages.get("Reference list", 12)),
        ("AI assistance disclosure", pages.get("AI assistance disclosure", 13)),
    ]
    table = doc.add_table(rows=len(entries), cols=2)
    set_table_geometry(table, [6.55, 0.63])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, (label, page_number) in enumerate(entries):
        left, right = table.rows[idx].cells
        for cell in (left, right):
            set_cell_margins(cell, 75, 40, 75, 40)
        lp = left.paragraphs[0]
        lp.paragraph_format.first_line_indent = None
        lp.paragraph_format.space_after = Pt(0)
        set_font(lp.add_run(label), "Carlito", 12)
        rp = right.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rp.paragraph_format.first_line_indent = None
        rp.paragraph_format.space_after = Pt(0)
        set_font(rp.add_run(str(page_number)), "Carlito", 12)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(14)
    set_bottom_border(p2, size=8)


def build():
    pages = json.loads(TOC_JSON.read_text()) if TOC_JSON.exists() else {}
    module, story = load_pdf_story()
    doc = Document()
    configure_styles(doc)
    doc.core_properties.title = "ITOPA3-33 - Assignment - Mowbray - EDUV4948467"
    doc.core_properties.author = "Joshua Nehohwa"
    doc.core_properties.subject = "Operating Systems individual assignment - AI-assisted study draft"

    cover_section = doc.sections[0]
    cover_section.page_width = Inches(8.5)
    cover_section.page_height = Inches(11)
    cover_section.top_margin = Inches(0.25)
    cover_section.bottom_margin = Inches(0.25)
    cover_section.left_margin = Inches(0.25)
    cover_section.right_margin = Inches(0.25)
    cover_p = doc.add_paragraph()
    cover_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_p.paragraph_format.space_before = Pt(0)
    cover_p.paragraph_format.space_after = Pt(0)
    cover_run = cover_p.add_run()
    cover_shape = cover_run.add_picture(str(COVER_IMAGE), width=Inches(8.0))
    cover_shape._inline.docPr.set("descr", "Eduvos Individual Assessment Coversheet")

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    body_section.page_width = Inches(8.5)
    body_section.page_height = Inches(11)
    body_section.top_margin = Inches(0.55)
    body_section.bottom_margin = Inches(0.58)
    body_section.left_margin = Inches(0.66)
    body_section.right_margin = Inches(0.66)
    body_section.header_distance = Inches(0.2)
    body_section.footer_distance = Inches(0.25)
    body_section.footer.is_linked_to_previous = False
    footer = body_section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(footer.add_run("ITOPA3-33_Assignment_Mowbray_eduv4948467 | Page "), "Arial", 9, color=RGBColor(0, 0, 0))
    add_page_field(footer)

    add_toc(doc, pages)
    doc.add_page_break()
    # The ReportLab story's first six items are the PDF table of contents and its page break.
    for flowable in story[6:]:
        process_flowable(doc, module, flowable)

    doc.settings.update_fields_on_open = True
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
