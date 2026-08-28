from copy import deepcopy
from pathlib import Path
from tempfile import NamedTemporaryFile
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


REFERENCE = Path('/Users/joshuanehohwa/Downloads/Eduvos Assessment Template.docx')
OUTPUT = Path('/Users/joshuanehohwa/Documents/Research and Projects/ITOPA Practical/ITOPA3-33 - Assignment - Mowbray - EDUV4948467 - Working Draft.docx')

NAVY = RGBColor(8, 45, 103)
GREY = RGBColor(89, 89, 89)
RED = RGBColor(156, 0, 6)


def set_cell_value(cell, value):
    paragraph = cell.paragraphs[0]
    paragraph.text = value
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    for run in paragraph.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(10.5)


def insert_before(anchor, text='', kind='body', bold_prefix=None):
    paragraph = anchor.insert_paragraph_before()
    if kind == 'page_break':
        paragraph.add_run().add_break(WD_BREAK.PAGE)
        return paragraph

    if kind == 'bullet':
        try:
            paragraph.style = 'List Bullet'
        except KeyError:
            pass
    elif kind == 'number':
        try:
            paragraph.style = 'List Number'
        except KeyError:
            pass

    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        first.bold = True
        paragraph.add_run(text[len(bold_prefix):])
    else:
        paragraph.add_run(text)

    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.08
    for run in paragraph.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(10.5)

    if kind == 'title':
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(10)
        for run in paragraph.runs:
            run.font.name = 'Georgia'
            run.font.size = Pt(18)
            run.font.color.rgb = NAVY
            run.bold = True
    elif kind == 'h1':
        paragraph.paragraph_format.space_before = Pt(11)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.keep_with_next = True
        for run in paragraph.runs:
            run.font.size = Pt(14)
            run.font.color.rgb = NAVY
            run.bold = True
    elif kind == 'h2':
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.keep_with_next = True
        for run in paragraph.runs:
            run.font.size = Pt(11.5)
            run.font.color.rgb = NAVY
            run.bold = True
    elif kind == 'note':
        for run in paragraph.runs:
            run.font.color.rgb = GREY
            run.italic = True
    elif kind == 'warning':
        paragraph.paragraph_format.keep_together = True
        for run in paragraph.runs:
            run.font.color.rgb = RED
            run.bold = True
    return paragraph


document = Document(REFERENCE)

# Coversheet metadata.
metadata = {
    (0, 0): 'Mowbray',
    (0, 1): 'Information Technology',
    (0, 2): 'ITOPA3-33',
    (0, 3): 'Bongani Mahlangu',
    (0, 4): 'N/A - Individual assignment',
    (1, 0): 'Assignment',
    (1, 2): 'Friday, 28 August 2026 at 23:59',
    (3, 0): 'Joshua Nehohwa',
    (3, 1): 'EDUV4948467',
}
for (table_index, row_index), value in metadata.items():
    set_cell_value(document.tables[table_index].rows[row_index].cells[1], value)

# Replace the assessment placeholder with a status line and insert the structured framework
# immediately before the AI Evidence heading, preserving the template's later sections.
document.paragraphs[12].text = 'Working response framework - Questions 1 and 2'
for run in document.paragraphs[12].runs:
    run.font.name = 'Georgia'
    run.font.size = Pt(18)
    run.font.color.rgb = NAVY
    run.bold = True

anchor = next(p for p in document.paragraphs if p.text.strip() == 'AI Evidence Section')

insert_before(anchor, 'Status and authorship note', 'h1')
insert_before(
    anchor,
    'This document contains source-backed planning notes and writing checkpoints. Convert every point into your own explanation, add your own examples and transitions, and verify it against the course material before submission.',
    'warning',
)
insert_before(anchor, 'Suggested writing target: approximately 500-650 words for Question 1 and 650-800 words for Question 2. The brief provides marks but no word count, so clarity and mark coverage take priority.', 'note')

insert_before(anchor, 'Question 1 - Operating-system system calls (20 marks)', 'h1')
insert_before(anchor, '1.1 File-management calls and their importance (8 marks)', 'h2')
insert_before(anchor, 'Writing task: explain the request as a controlled sequence from the banking application to the OS, not merely as a list of function names.', 'note')
for text in [
    'Locate/open: the application supplies the statement path; the OS resolves the directory path, checks the customer/process permissions, creates an open-file description and returns a protected handle or file descriptor.',
    'Read: the application requests statement bytes through the descriptor; the OS coordinates the file system, cache/buffers, storage driver and hardware, then copies authorised data into application memory.',
    'Seek or metadata lookup where needed: the OS can obtain file size, timestamps or a required offset before reading the relevant statement content.',
    'Write/spool only if the statement is exported or printed: the OS controls output buffering, printer queues and device-driver communication.',
    'Close: the descriptor/handle is released after display or printing, allowing the OS to reclaim per-process resources and complete pending I/O safely.',
    'Importance to emphasise: abstraction from device-specific commands, authentication and access control, consistent error reporting, safe sharing/concurrency, auditing, caching and reliable resource cleanup.',
]:
    insert_before(anchor, text, 'bullet')
insert_before(anchor, 'Course anchors: Week 1 slides 41, 43, 49 and 59-60; Week 2 slide 27. External verification: POSIX open(), read() and close() interfaces.', 'note')
insert_before(anchor, 'Student checkpoint: add one sentence showing how each call protects the confidentiality and integrity of a customer statement.', 'warning')

insert_before(anchor, '1.2 Consequences of bypassing the OS (6 marks)', 'h2')
insert_before(anchor, 'Build the analysis around consequences and trade-offs. Acknowledge the proposed performance motive, then test whether it outweighs the following risks:', 'body')
for text in [
    'Security: direct device access bypasses normal permission checks, protected handles, process isolation and audit controls, exposing other customers\' data and creating an escalation path.',
    'Integrity and consistency: raw disk writes can ignore file-system metadata, journaling, locks and cache-coherency rules, corrupting statements, transaction records or the whole volume.',
    'Reliability: an application defect could address the wrong block/device, hang the printer, leak privileged resources or crash system-level components.',
    'Concurrency: several banking sessions could issue conflicting device operations without OS scheduling, locking, queuing and arbitration.',
    'Portability and maintainability: the application becomes tied to a particular disk, controller and printer protocol instead of stable OS interfaces and drivers.',
    'Performance reality: any saved call overhead may be outweighed by lost caching, buffering, DMA, disk scheduling and printer spooling; the proposal also moves complex driver logic into the banking application.',
]:
    insert_before(anchor, text, 'bullet')
insert_before(anchor, 'Course anchors: Week 1 slides 41, 43, 49 and 59-60. External verification: Microsoft Windows security model and I/O-manager guidance.', 'note')

insert_before(anchor, '1.3 CIO decision and justification (6 marks)', 'h2')
insert_before(anchor, 'Decision checkpoint: state an unambiguous rejection of direct hardware access in 1.3.1, then use 1.3.2 to justify the decision through the OS trust boundary.', 'warning')
for text in [
    'System calls are the controlled gateway from user mode into privileged kernel services.',
    'They enforce protection, validate requests, coordinate shared resources, expose stable interfaces and return defined errors.',
    'In a regulated banking environment, confidentiality, correctness, auditability and recoverability outweigh a speculative micro-optimisation.',
    'If performance is inadequate, optimise within supported mechanisms: profiling, asynchronous I/O, caching, batching, storage upgrades or approved driver/platform changes.',
]:
    insert_before(anchor, text, 'bullet')

insert_before(anchor, 'Question 2 - Multithreaded ShopFast web server (25 marks)', 'h1')
insert_before(anchor, '2.1 Single-threaded performance and blocking (6 marks)', 'h2')
for text in [
    '2.1a: explain serial service - one request must complete before the next starts. During Black Friday, arrival rate exceeds service rate, the waiting queue grows, latency rises, requests time out and the four-core machine remains underused.',
    '2.1b: define blocking as a thread being unable to progress while waiting for I/O or another event. A database lookup, payment API call or network write can therefore stop the only server thread, preventing unrelated customers from being served.',
    'Keep throughput and responsiveness distinct: throughput is completed requests per unit time; responsiveness is how quickly an individual customer receives progress or a response.',
]:
    insert_before(anchor, text, 'bullet')
insert_before(anchor, 'Course anchors: Week 2 slides 20-21 and 36 for blocking; slides 44-46 for threads, responsiveness and multicore execution.', 'note')

insert_before(anchor, '2.2 Three scenario-specific benefits (6 marks)', 'h2')
insert_before(anchor, 'Use exactly three benefits and pair every benefit with a ShopFast consequence:', 'body')
for text in [
    'Responsiveness/concurrency: while one request waits for a database, payment gateway or network operation, other worker threads can continue serving product and cart requests, reducing timeouts.',
    'Economy/resource sharing: threads share the server process\'s code, configuration, caches and address space, with lower creation/context-switch overhead than separate processes.',
    'Scalability/parallelism: runnable request handlers can execute on different cores of the quad-core server, increasing throughput during peaks when work and shared-resource contention are properly managed.',
]:
    insert_before(anchor, text, 'bullet')
insert_before(anchor, 'Course anchor: Week 2 slide 46 lists responsiveness, resource sharing, economy and scalability.', 'note')

insert_before(anchor, '2.3 Task parallelism versus data parallelism (6 marks)', 'h2')
for text in [
    'Data parallelism: partition a dataset across cores and apply the same operation to each subset.',
    'Task parallelism: distribute different tasks across cores; tasks may perform different operations on the same or different data.',
    'Primary classification for this scenario: task parallelism. Separate request tasks - browsing, cart updates and payments - can run at the same time and can follow different code paths.',
    'Two reasons to develop: requests are independent units belonging to different customers, and the handlers can execute concurrently across four cores even when their operations/data differ.',
    'Nuance: repeated instances of the same handler over many customer records resemble data parallelism, but the scenario\'s overall web-server architecture is best justified as task parallelism.',
]:
    insert_before(anchor, text, 'bullet')
insert_before(anchor, 'Course anchor: Week 2 slide 47.', 'note')

insert_before(anchor, '2.4 Thread pool and task queue (7 marks)', 'h2')
for text in [
    '2.4a: define a thread pool as a managed collection of reusable worker threads created ahead of, or independently from, individual requests.',
    '2.4b: explain why pooling is preferred - it avoids repeated creation/destruction overhead, bounds the number of active threads, reduces scheduler/memory pressure and prevents a traffic spike from creating unlimited threads.',
    '2.4c: describe the task queue as the waiting area between request acceptance and execution. Incoming requests become tasks; an available worker removes a task, processes it and returns to the pool.',
    'Queue role to connect to ShopFast: absorb short bursts, enforce ordering or priority, provide backpressure through bounded capacity, and expose overload when both workers and queue are saturated.',
    'Quad-core nuance: do not claim that exactly four workers is always optimal. CPU-bound work may be close to core count, while I/O-bound requests may benefit from more workers; the pool must still be bounded and measured.',
]:
    insert_before(anchor, text, 'bullet')
insert_before(anchor, 'Course anchor: Week 2 slide 54. External verification: Oracle Java ThreadPoolExecutor documents pooled workers, bounded resource management and the work queue.', 'note')

insert_before(anchor, 'Q1-Q2 self-review checklist', 'h1')
for text in [
    'Every sub-question is visibly labelled and answered in proportion to its marks.',
    'Every technical definition is followed by application to the named scenario.',
    'The final prose distinguishes an OS service, a system call and a device driver.',
    'The final prose does not promise that more threads always improve performance; it acknowledges bounded resources and contention.',
    'All wording has been rewritten in your own voice, with your own transitions and examples.',
    'In-text citations correspond to entries in the working reference list.',
]:
    insert_before(anchor, text, 'bullet')

insert_before(anchor, 'Working reference list', 'h1')
references = [
    'Eduvos. 2026a. ITOPA3-T11 Week 0 & 1 Lessons 1-4. Internal course slides.',
    'Eduvos. 2026b. ITOPA3-T11 Week 2 Lessons 5-7. Internal course slides.',
    'The Open Group. 2024. POSIX.1-2024 System Interfaces: open, read and close. Available at: https://pubs.opengroup.org/onlinepubs/9799919799/idx/functions.html',
    'Microsoft. 2025. Windows security model for driver developers. Available at: https://learn.microsoft.com/en-us/windows-hardware/drivers/driversecurity/windows-security-model',
    'Microsoft. 2025. User mode and kernel mode. Available at: https://learn.microsoft.com/en-us/windows-hardware/drivers/gettingstarted/user-mode-and-kernel-mode',
    'Oracle. 2025. ThreadPoolExecutor, Java SE 24 API. Available at: https://docs.oracle.com/en/java/javase/24/docs/api/java.base/java/util/concurrent/ThreadPoolExecutor.html',
]
for text in references:
    insert_before(anchor, text, 'body')

insert_before(anchor, 'Questions 3 and 4 - next work block', 'h1')
insert_before(anchor, 'Question 3 is supported by the supplied Week 3 deck (deadlock system model, resource-allocation graphs, prevention and avoidance). Question 4 requires the assignment diagram plus Week 5-6 file-management/cache material, which has not yet been supplied.', 'note')

# AI evidence: record the actual prompts and the actions taken, without inventing a policy code.
evidence = next(p for p in document.paragraphs if p.text.strip().startswith('[AI Code 1'))
evidence.text = 'Prompt record 1'
evidence.paragraph_format.space_before = Pt(8)
evidence.paragraph_format.space_after = Pt(4)
evidence.paragraph_format.keep_with_next = True
for run in evidence.runs:
    run.font.name = 'Arial'
    run.font.size = Pt(11.5)
    run.font.color.rgb = NAVY
    run.bold = True

evidence_anchor = next(p for p in document.paragraphs if p.text.strip() == 'AI Disclosure Appendix')
anchor.paragraph_format.page_break_before = True
evidence_position = next(index for index, paragraph in enumerate(document.paragraphs) if paragraph._p is evidence._p)
prompt_anchor = document.paragraphs[evidence_position + 1]
insert_before(prompt_anchor, '"If you go to my downloads on my macbook, there is a file title ITOPA practical, its recent, extract it here please and go through it and let me know what it needs me to do."', 'body')
insert_before(prompt_anchor, 'Action: located and copied the ITOPA assignment brief, extracted its eight pages, and summarised the questions, marks and submission rules.', 'note')
insert_before(prompt_anchor, 'Prompt record 2', 'h2')
insert_before(prompt_anchor, '"Its due Friday night, so lets plan and get it done."', 'body')
insert_before(prompt_anchor, 'Action: created a mark-weighted schedule and authorship workflow around the Friday 23:59 deadline.', 'note')
insert_before(prompt_anchor, 'Prompt record 3', 'h2')
insert_before(prompt_anchor, '"mowbray, eduv4948467, Bongani Mahlangu, use these files for now" with three Eduvos ITOPA slide decks supplied.', 'body')
insert_before(prompt_anchor, 'Action: inspected the supplied Week 0-3 decks, mapped exact slides to Questions 1 and 2, checked gaps against official POSIX, Microsoft and Oracle documentation, and produced a structured working framework rather than final submission prose.', 'note')

# AI disclosure table: replace example content with a truthful single-tool entry.
disclosure = document.tables[6]
values = [
    'OpenAI Codex (GPT-5.6)',
    'Located and interpreted the brief; analysed Week 0-3 slides; developed a Q1-Q2 research and writing framework; checked technical claims; formatted the working document.',
    'Mapped every point to the supplied Eduvos slides and official POSIX, Microsoft and Oracle documentation. Final assessed wording, examples and reasoning must be rewritten and checked by the student.',
]
for column, value in enumerate(values):
    set_cell_value(disclosure.rows[1].cells[column], value)
for row in (2, 3):
    for cell in disclosure.rows[row].cells:
        set_cell_value(cell, '')

# This is an individual assignment, so retain one signature row instead of the
# six-row group signature grid. The smaller grid also prevents a trailing blank page.
signature_table = document.tables[7]
while len(signature_table.rows) > 2:
    signature_table._tbl.remove(signature_table.rows[-1]._tr)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
document.save(OUTPUT)

# python-docx can rewrite unchanged package parts. Restore every reference part
# except the intentionally edited main document XML so the template's styles,
# relationships, headers, footers, theme and custom XML remain byte-identical.
with ZipFile(REFERENCE, 'r') as source_zip, ZipFile(OUTPUT, 'r') as edited_zip:
    edited_document_xml = edited_zip.read('word/document.xml')
    output_entries = edited_zip.namelist()
    source_entries = set(source_zip.namelist())
    with NamedTemporaryFile(suffix='.docx', delete=False, dir=OUTPUT.parent) as temp_file:
        temp_path = Path(temp_file.name)
    with ZipFile(temp_path, 'w', ZIP_DEFLATED) as final_zip:
        for name in output_entries:
            if name == 'word/document.xml':
                payload = edited_document_xml
            elif name in source_entries:
                payload = source_zip.read(name)
            else:
                payload = edited_zip.read(name)
            final_zip.writestr(name, payload)
temp_path.replace(OUTPUT)
print(OUTPUT)
