from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_assignment_pack import PAPERS, add_page_number_footer, configure_styles


OUT_DIR = Path(__file__).resolve().parent


def add_placeholder(paragraph, text: str) -> None:
    run = paragraph.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(100, 116, 139)


def add_body_placeholder(doc: Document, prompt: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Inches(0.25)
    add_placeholder(paragraph, prompt)


def add_citation_anchor(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run("Citation anchors: ")
    run.bold = True
    run.font.color.rgb = RGBColor(31, 77, 120)
    add_placeholder(paragraph, text)


def create_submission_template() -> Path:
    doc = Document()
    configure_styles(doc)
    add_page_number_footer(doc)

    title = doc.add_paragraph()
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(
        "Risk Mitigation in CI/CD Pipelines for Safety-Critical Embedded Systems Using Automated Testing and Rollback Strategies"
    ).bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("ITDMA3-22 Research Design and Methodology | Assignment 1").italic = True

    details = [
        "Student name: [Insert your name]",
        "Student number: [Insert your student number]",
        "Campus: [Insert campus name]",
        "Submission date: 8 June 2026",
        "Referencing style: Harvard",
    ]
    for detail in details:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_placeholder(paragraph, detail)

    doc.add_paragraph()

    doc.add_heading("Abstract", level=1)
    add_body_placeholder(
        doc,
        "[Write 150-200 words. Briefly state the problem, the literature review method, the main synthesis, and the recommended risk-aware CI/CD pipeline approach.]",
    )

    doc.add_heading("Keywords", level=1)
    add_body_placeholder(
        doc,
        "[continuous integration; continuous deployment; safety-critical embedded systems; automated testing; rollback; graceful degradation]",
    )

    doc.add_heading("1. Introduction", level=1)
    add_body_placeholder(
        doc,
        "[Write 350-450 words. Introduce the tension between faster software delivery and safety-critical reliability. Explain why embedded systems differ from web software because they involve hardware, timing, certification, field validation, and failure consequences. End with the research objectives and the research question.]",
    )
    add_citation_anchor(
        doc,
        "Use Zampetti et al. (2023), Dakkak et al. (2023), Baron and Louis (2023).",
    )

    doc.add_heading("2. Methodology", level=1)
    add_body_placeholder(
        doc,
        "[Write 450-600 words. Explain that this article is based on a structured literature review of eight peer-reviewed papers. Name the databases searched, search terms, inclusion/exclusion criteria, and comparison fields: section structure, methodology, research approach, framework/theory, data collection, sampling, data analysis, findings, contradictions, gaps, and relevance.]",
    )
    add_citation_anchor(
        doc,
        "Use Garousi et al. (2018) as an example of systematic mapping logic if needed.",
    )

    doc.add_heading("3. Literature Review", level=1)
    add_body_placeholder(
        doc,
        "[Write 650-800 words. Organise this section by themes rather than summarising one paper at a time. Cover CI/CD in cyber-physical systems, continuous deployment in systems-of-systems, embedded automated testing, mutation testing, digital twin testing, continuous certification, and graceful degradation.]",
    )
    add_citation_anchor(
        doc,
        "Use Zampetti et al. (2023), Dakkak et al. (2023), Garousi et al. (2018), Du et al. (2022), Baker and Habli (2013), Barbie and Hasselbring (2024), Baron and Louis (2023), Becker, Voss and Schätz (2018).",
    )

    doc.add_heading("4. Results and Comparative Analysis", level=1)
    add_body_placeholder(
        doc,
        "[Write 800-1000 words. Compare how the selected papers present their sections. Then compare methodologies: interviews/survey, industrial case study, engineering prototype, empirical mutation-testing evaluation, systematic mapping, conceptual modelling, certification framework, and formal modelling. Identify research approaches, data collection, sampling and data analysis methods.]",
    )
    add_body_placeholder(
        doc,
        "[Make sure this section explicitly identifies findings that confirm prior studies, extend existing research, contradict or qualify earlier work, and show agreement across the literature.]",
    )

    doc.add_heading("5. Discussion", level=1)
    add_body_placeholder(
        doc,
        "[Write 600-800 words. Answer the research question directly. Evaluate what the findings mean for your proposed risk-aware CI/CD pipeline. Discuss why automated testing reduces risk before deployment, and why rollback/fail-safe strategies reduce harm after faults escape.]",
    )
    add_body_placeholder(
        doc,
        "[Discuss literature gaps: limited integrated studies connecting CI/CD, safety certification, digital twins/HIL, mutation testing, staged deployment and rollback in one end-to-end pipeline; limited evidence across regulated domains; and limited clarity on when rollback is safer than graceful degradation.]",
    )
    add_citation_anchor(
        doc,
        "Use Becker, Voss and Schätz (2018) for graceful degradation; Baker and Habli (2013) for test adequacy; Baron and Louis (2023) for continuous certification.",
    )

    doc.add_heading("6. Conclusion and Recommendations", level=1)
    add_body_placeholder(
        doc,
        "[Write 250-350 words. Summarise the answer to the research question. Recommend a layered pipeline: build/static checks, simulation or digital twin tests, HIL tests, safety-assurance evidence, staged deployment, monitoring, and rollback/fail-safe/degraded-mode recovery.]",
    )

    doc.add_heading("7. Proposed Journal or Conference", level=1)
    add_body_placeholder(
        doc,
        "[Write 100-150 words. Identify the International Conference on Software Testing, Verification and Validation (ICST) as a suitable venue because it covers software testing, verification, validation, empirical studies, tools, embedded software and technology transfer.]",
    )

    doc.add_heading("References", level=1)
    for paper in sorted(PAPERS, key=lambda item: item["reference"].casefold()):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.add_run(paper["reference"])

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Arial"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
            if run.font.size is None:
                run.font.size = Pt(12)

    out = OUT_DIR / "ITDMA3-22_Assignment_1_Submission_Template_Harvard.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    print(create_submission_template())
