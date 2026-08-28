from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


OUT_DIR = Path(__file__).resolve().parent


PAPERS = [
    {
        "id": "P1",
        "citation": "Zampetti et al. (2023)",
        "in_text": "(Zampetti et al., 2023)",
        "reference": "Zampetti, F., Tamburri, D.A., Panichella, S., Panichella, A., Canfora, G. and Di Penta, M. (2023) 'Continuous integration and delivery practices for cyber-physical systems: An interview-based study', ACM Transactions on Software Engineering and Methodology, 32(3), pp. 1-44. Available at: https://doi.org/10.1145/3571854.",
        "url": "https://research.tudelft.nl/en/publications/continuous-integration-and-delivery-practices-for-cyber-physical-",
        "peer_review": "Peer-reviewed journal article in ACM TOSEM.",
        "purpose": "Empirically investigates CI/CD challenges, barriers and mitigation practices for cyber-physical systems across ten organisations.",
        "structure": "Standard empirical article: abstract, introduction, related work/background, research method, findings, discussion/implications, threats to validity and conclusion.",
        "methodology": "Interview-based empirical software engineering study using semi-structured interviews, open card sorting, member checking and a validating survey with 55 professional developers.",
        "approach": "Interpretivist qualitative core with confirmatory survey validation; pragmatic/mixed empirical orientation.",
        "framework": "CI/CD practices for CPS; empirical software engineering coding and validation procedures.",
        "data_collection": "Semi-structured interviews in 10 organisations and follow-up survey.",
        "sampling": "Purposive organisational sampling across eight CPS domains; professional developer survey sample.",
        "analysis": "Open card sorting, qualitative categorisation, member checking and survey validation.",
        "findings": "CPS CI/CD requires a balance of continuous and periodic builds, careful use of simulators and Hardware-in-the-Loop, attention to deployment difficulty, and teams that combine hardware and software expertise.",
        "confirms": "Confirms that CI/CD improves feedback and defect discovery, but embedded/CPS contexts add hardware and deployment constraints.",
        "extends": "Extends CI/CD literature by showing CPS-specific barriers around simulator-HIL mismatch and interdisciplinary expertise.",
        "contradicts": "Challenges the assumption that web-style continuous deployment practices transfer directly into CPS environments.",
        "gaps": "Limited direct design of risk-aware rollback mechanisms; mostly organisational and practice-level evidence.",
        "use_in_article": "Use as the main anchor for why safety-critical embedded CI/CD must combine automation with periodic assurance gates.",
    },
    {
        "id": "P2",
        "citation": "Dakkak et al. (2023)",
        "in_text": "(Dakkak et al., 2023)",
        "reference": "Dakkak, A., Bosch, J., Holmström Olsson, H. and Mattos, D.I. (2023) 'Continuous deployment in software-intensive system-of-systems', Information and Software Technology, 159, Article 107200. Available at: https://doi.org/10.1016/j.infsof.2023.107200.",
        "url": "https://research.tue.nl/en/publications/continuous-deployment-in-software-intensive-system-of-systems/",
        "peer_review": "Peer-reviewed journal article in Information and Software Technology.",
        "purpose": "Studies how software-intensive systems-of-systems transition to continuous deployment, especially field testing and validation.",
        "structure": "Empirical article with context, objectives, method, results, discussion and conclusion.",
        "methodology": "Industrial case study at Ericsson AB focused on 3G Radio Access Network embedded software.",
        "approach": "Pragmatic/interpretive case study using industrial evidence.",
        "framework": "Continuous deployment transition and orchestration of field testing/validation in systems-of-systems.",
        "data_collection": "Industrial case data, interviews and organisational deployment/validation evidence described by the authors.",
        "sampling": "Single reference case: Ericsson 3G RAN system-of-systems.",
        "analysis": "Case-study analysis of deployment cadence, validation orchestration and success factors.",
        "findings": "Deployment and field validation become continuous; orchestration between constituent systems, monitoring, documentation and management support become essential.",
        "confirms": "Confirms the need for testing, monitoring and organisational readiness in continuous deployment.",
        "extends": "Extends embedded deployment research from single systems to software-intensive systems-of-systems.",
        "contradicts": "Contradicts simplistic 'deploy fast and rollback later' thinking by showing that field validation must be planned continuously.",
        "gaps": "Not a safety-critical certification study; rollback is discussed less directly than monitoring and validation.",
        "use_in_article": "Use to explain staged deployment, field validation and monitoring as risk controls in complex embedded environments.",
    },
    {
        "id": "P3",
        "citation": "Du et al. (2022)",
        "in_text": "(Du et al., 2022)",
        "reference": "Du, B., Azimi, S., Moramarco, A., Sabena, D., Parisi, F. and Sterpone, L. (2022) 'An automated continuous integration multitest platform for automotive systems', IEEE Systems Journal, 16(2), pp. 2495-2506. Available at: https://doi.org/10.1109/JSYST.2021.3069548.",
        "url": "https://iris.polito.it/handle/11583/2898672",
        "peer_review": "Peer-reviewed IEEE Systems Journal article.",
        "purpose": "Presents a CI testing framework using virtual and FPGA-based verification platforms for automotive systems.",
        "structure": "Engineering design/evaluation article: introduction, background, platform design, experimental demonstration, results and conclusion.",
        "methodology": "Design science / engineering prototype evaluated on a real heterogeneous automotive system.",
        "approach": "Pragmatic positivist engineering evaluation.",
        "framework": "Continuous integration pipeline with commercial virtual platform and FPGA verification.",
        "data_collection": "Automated test runs and platform performance/validation observations.",
        "sampling": "Automotive engine-control-unit focused case/platform demonstration.",
        "analysis": "Technical comparison of test execution feasibility, efficiency and viability.",
        "findings": "Multi-stage CI testing can reduce late defect discovery and support hardware-software co-design before final hardware is fully available.",
        "confirms": "Confirms the value of earlier and automated testing in complex automotive systems.",
        "extends": "Extends CI into automotive verification by combining virtual and hardware-based test stages.",
        "contradicts": "Shows that pure software-only CI is insufficient where hardware behavior matters.",
        "gaps": "Limited explicit treatment of rollback or certification evidence; focused on testing platform feasibility.",
        "use_in_article": "Use as concrete evidence for automated testing techniques suitable for embedded CI/CD pipelines.",
    },
    {
        "id": "P4",
        "citation": "Baker and Habli (2013)",
        "in_text": "(Baker and Habli, 2013)",
        "reference": "Baker, R. and Habli, I. (2013) 'An empirical evaluation of mutation testing for improving the test quality of safety-critical software', IEEE Transactions on Software Engineering, 39(6), pp. 787-805. Available at: https://doi.org/10.1109/TSE.2012.56.",
        "url": "https://pure.york.ac.uk/portal/en/publications/an-empirical-evaluation-of-mutation-testing-for-improving-the-tes/",
        "peer_review": "Peer-reviewed IEEE Transactions on Software Engineering article.",
        "purpose": "Evaluates whether mutation testing can improve test quality for airborne safety-critical software that already satisfied coverage requirements.",
        "structure": "Empirical evaluation article with abstract, introduction, background, study design, results, analysis, industry feedback and conclusion.",
        "methodology": "Empirical evaluation of mutation testing on safety-critical C and Ada airborne software.",
        "approach": "Positivist quantitative/empirical evaluation with industry feedback.",
        "framework": "Mutation testing, structural coverage and safety-critical verification lifecycle.",
        "data_collection": "Mutation results, test failures and industry feedback.",
        "sampling": "Safety-critical airborne software programs written in high-integrity subsets of C and Ada.",
        "analysis": "Mutation survival analysis, root-cause analysis and relationship between program characteristics and mutation results.",
        "findings": "Mutation testing can reveal weaknesses missed by traditional coverage and manual review, and can point to issues in requirements and coding.",
        "confirms": "Confirms that testing is central to certification assurance.",
        "extends": "Extends assurance thinking beyond structural coverage by adding mutation-based sufficiency evidence.",
        "contradicts": "Challenges the view that mandated coverage alone proves test adequacy.",
        "gaps": "Older study and not CI/CD-specific; needs linking to automated pipeline gates.",
        "use_in_article": "Use to argue that safety-critical CI pipelines should include stronger test adequacy checks, not just pass/fail unit tests.",
    },
    {
        "id": "P5",
        "citation": "Garousi et al. (2018)",
        "in_text": "(Garousi et al., 2018)",
        "reference": "Garousi, V., Felderer, M., Karapıçak, Ç.M. and Yılmaz, U. (2018) 'Testing embedded software: A survey of the literature', Information and Software Technology, 104, pp. 14-45. Available at: https://doi.org/10.1016/j.infsof.2018.06.016.",
        "url": "https://research.wur.nl/en/publications/testing-embedded-software-a-survey-of-the-literature",
        "peer_review": "Peer-reviewed Information and Software Technology article.",
        "purpose": "Maps the embedded software testing literature and classifies techniques, artefacts and industrial contexts.",
        "structure": "Systematic mapping/survey article: context, objective, method, search and selection, mapping plan, results, discussion, threats and conclusion.",
        "methodology": "Systematic literature review / systematic literature mapping.",
        "approach": "Positivist evidence synthesis with structured classification.",
        "framework": "Systematic mapping classification of embedded testing topics, activities, artefacts and industries.",
        "data_collection": "Initial pool of 588 papers, final pool of 312 technical papers, plus practitioner feedback.",
        "sampling": "Literature selected through explicit inclusion/exclusion criteria.",
        "analysis": "Systematic classification, frequency mapping and synthesis.",
        "findings": "Embedded testing is broad and fragmented; practitioners need structured guidance to avoid reinventing testing approaches.",
        "confirms": "Confirms the importance and diversity of embedded testing techniques.",
        "extends": "Extends single-technique studies by providing a wide testing landscape.",
        "contradicts": "Does not contradict directly; instead exposes fragmentation and weak synthesis in the field.",
        "gaps": "Survey is broad; it does not prescribe a specific risk-aware CI/CD architecture.",
        "use_in_article": "Use as the broad evidence base for selecting automated testing layers in the pipeline.",
    },
    {
        "id": "P6",
        "citation": "Barbie and Hasselbring (2024)",
        "in_text": "(Barbie and Hasselbring, 2024)",
        "reference": "Barbie, A. and Hasselbring, W. (2024) 'From digital twins to digital twin prototypes: Concepts, formalization, and applications', IEEE Access, 12, pp. 75337-75365. Available at: https://doi.org/10.1109/ACCESS.2024.3406510.",
        "url": "https://www.cau-se.de/news/2024-06new_journal_paper_on_digital_twin_formalization/",
        "peer_review": "Peer-reviewed IEEE Access article.",
        "purpose": "Formalises digital twin prototype concepts and shows how they support automated testing of embedded software.",
        "structure": "Conceptual and applied article: introduction, concept formalisation, models, applications/field studies, replication material and conclusion.",
        "methodology": "Conceptual formalisation using Object-Z and UML, with application in real-world field studies and a lab replication study.",
        "approach": "Design science and conceptual modelling.",
        "framework": "Digital twin prototype concepts, digital thread/shadow/twin model and automated integration testing.",
        "data_collection": "Field-study and laboratory application evidence.",
        "sampling": "Two field studies plus a published lab study.",
        "analysis": "Conceptual modelling and applied demonstration.",
        "findings": "Digital twin prototypes allow virtual-context automated integration tests in CI/CD without constant access to the physical object.",
        "confirms": "Confirms that hardware dependency is a bottleneck in embedded CI testing.",
        "extends": "Extends simulator/HIL discussions by proposing digital twin prototypes as a CI-friendly testing substitute.",
        "contradicts": "Qualifies reliance on physical HIL by showing a virtual test layer can reduce access constraints.",
        "gaps": "Digital twin fidelity, safety certification acceptance and rollback integration remain open questions.",
        "use_in_article": "Use for simulation and digital-twin testing as a scalable complement to HIL.",
    },
    {
        "id": "P7",
        "citation": "Baron and Louis (2023)",
        "in_text": "(Baron and Louis, 2023)",
        "reference": "Baron, C. and Louis, V. (2023) 'Framework and tooling proposals for Agile certification of safety-critical embedded software in avionic systems', Computers in Industry, 148, Article 103887. Available at: https://doi.org/10.1016/j.compind.2023.103887.",
        "url": "https://www.sciencedirect.com/science/article/abs/pii/S0166361523000374",
        "peer_review": "Peer-reviewed Computers in Industry article.",
        "purpose": "Proposes a framework and tooling for continuous certification of DO-178C/ED-12C safety-critical avionics software.",
        "structure": "Framework proposal article: introduction, certification problem, framework/tooling proposal, compliance discussion, industrial case and conclusion.",
        "methodology": "Framework/tooling proposal illustrated through an industrial case study.",
        "approach": "Pragmatic design science with compliance-oriented evaluation.",
        "framework": "Continuous certification, Agile-compatible certification process and DO-178C/ED-12C compliance thinking.",
        "data_collection": "Industrial case evidence and certification-process analysis.",
        "sampling": "Avionics safety-critical software certification context.",
        "analysis": "Framework compliance analysis and industrial case evaluation.",
        "findings": "Certification requirements should be integrated continuously into the development lifecycle instead of treated as late audit work.",
        "confirms": "Confirms that safety-critical software needs auditable evidence and regulatory alignment.",
        "extends": "Extends Agile/CI practice into certification by proposing continuous certification tooling.",
        "contradicts": "Challenges the traditional V-cycle habit of deferring certification evidence until late stages.",
        "gaps": "Avionics-specific; needs careful transfer to broader embedded domains like automotive or medical systems.",
        "use_in_article": "Use to support the argument that the pipeline must generate traceable assurance artefacts, not merely deploy software faster.",
    },
    {
        "id": "P8",
        "citation": "Becker, Voss and Schätz (2018)",
        "in_text": "(Becker, Voss and Schätz, 2018)",
        "reference": "Becker, K., Voss, S. and Schätz, B. (2018) 'Formal analysis of feature degradation in fault-tolerant automotive systems', Science of Computer Programming, 154, pp. 89-133. Available at: https://doi.org/10.1016/j.scico.2017.10.007.",
        "url": "https://www.sciencedirect.com/science/article/pii/S0167642317302198",
        "peer_review": "Peer-reviewed Science of Computer Programming article.",
        "purpose": "Analyses graceful degradation scenarios for mixed-criticality automotive systems.",
        "structure": "Formal methods article: introduction, system model, formal analysis, case/application, related work and conclusion.",
        "methodology": "Formal modelling and constraint-based analysis using an SMT solver, applied to a constructed automotive example.",
        "approach": "Positivist/formal analytical approach.",
        "framework": "Graceful degradation, redundancy, fail-operational and mixed-criticality system modelling.",
        "data_collection": "Formal model inputs and constructed automotive scenario.",
        "sampling": "Constructed automotive example, not empirical population sampling.",
        "analysis": "Constraint solving and Boolean availability/degradation analysis.",
        "findings": "Safety-critical embedded systems need failure-handling actions such as fail-silent, fail-safe, fail-operational and degraded modes to avoid harmful system-level failure.",
        "confirms": "Confirms that deployment and recovery decisions must be designed before runtime failure happens.",
        "extends": "Extends rollback thinking from simple version reversal to controlled feature degradation and redundancy-aware deployment.",
        "contradicts": "Challenges simplistic rollback-only recovery for mixed-critical systems where safe degraded operation may be preferable.",
        "gaps": "Formal model is not directly integrated with a CI/CD pipeline.",
        "use_in_article": "Use as the main theoretical basis for rollback, fail-safe and graceful degradation mechanisms.",
    },
]


SECTION_GUIDE = [
    {
        "section": "Abstract",
        "purpose": "Summarises the topic, method, key synthesis and contribution.",
        "content": "Problem context, literature-search method, main findings, gaps and recommendation.",
        "style": "Concise, objective and written after the article is complete.",
        "citation_hint": "Usually no citations unless your lecturer expects them; focus on your synthesis.",
    },
    {
        "section": "Introduction",
        "purpose": "Frames why CI/CD risk mitigation matters in safety-critical embedded systems.",
        "content": "Context, problem, objectives, research question and article roadmap.",
        "style": "Persuasive but evidence-based; move from broad problem to focused question.",
        "citation_hint": "Use sources such as Zampetti et al. (2023), Dakkak et al. (2023) and Baron and Louis (2023).",
    },
    {
        "section": "Literature Review",
        "purpose": "Synthesises existing knowledge before presenting your analysis.",
        "content": "CI/CD in CPS, embedded testing, mutation testing, digital twins, certification and graceful degradation.",
        "style": "Analytical, theme-based and comparative rather than source-by-source summaries.",
        "citation_hint": "Every major claim should be attached to a Harvard citation.",
    },
    {
        "section": "Methodology",
        "purpose": "Explains how you found, selected and analysed the 8 papers.",
        "content": "Databases, search terms, inclusion/exclusion criteria and comparison matrix fields.",
        "style": "Objective and reproducible; avoid overstating this as a full systematic review.",
        "citation_hint": "Cite Garousi et al. (2018) as an example of systematic mapping logic if useful.",
    },
    {
        "section": "Results / Comparative Analysis",
        "purpose": "Shows what the papers contain and how their structures/methods differ.",
        "content": "Section comparison, research approaches, data collection, sampling, analysis methods and findings.",
        "style": "Comparative and evidence-driven; use grouped patterns.",
        "citation_hint": "Use clusters of citations, e.g. (Zampetti et al., 2023; Dakkak et al., 2023).",
    },
    {
        "section": "Discussion",
        "purpose": "Evaluates what the findings mean for a risk-aware CI/CD pipeline.",
        "content": "Agreement, contradictions, extensions, gaps, implications and pipeline recommendations.",
        "style": "Most analytical section; link back to research question and objectives.",
        "citation_hint": "Use citations while making your own evaluation visible.",
    },
    {
        "section": "Conclusion",
        "purpose": "Answers the research question and closes the argument.",
        "content": "Main answer, contribution, limitations and future research direction.",
        "style": "Clear, restrained and no new evidence.",
        "citation_hint": "Citations are optional; avoid introducing new sources here.",
    },
]


PIPELINE_MODEL = [
    "Commit and build gate: static checks, unit tests and traceability checks.",
    "Simulation / digital twin gate: fast repeatable integration tests using simulators or digital twin prototypes.",
    "Hardware-in-the-Loop gate: selected tests on real or FPGA-supported hardware for timing and physical interaction risks.",
    "Safety assurance gate: mutation/coverage evidence, requirements traceability and certification artefact checks.",
    "Staged deployment gate: limited rollout, monitoring and field validation before wider release.",
    "Rollback / fail-safe gate: version rollback, feature degradation, fail-silent/fail-safe/fail-operational mode selection.",
]


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D0D7DE")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.style = doc.styles["Footer"]
    run = footer.add_run("ITDMA Assignment 1 Working Pack | Page ")
    run.font.size = Pt(9)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(8)

    for style_name, size, color, before, after in [
        ("Title", 20, "111827", 0, 10),
        ("Heading 1", 16, "1F4D78", 14, 6),
        ("Heading 2", 13, "2E74B5", 10, 5),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(4)


def add_doc_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    header_cells = table.rows[0].cells
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = header_cells[idx]
        shade_cell(cell, "E8EEF5")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            run.font.name = "Arial"
            run.font.size = Pt(8.5)
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
    doc.add_paragraph()


def create_docx() -> Path:
    doc = Document()
    configure_styles(doc)
    add_page_number_footer(doc)

    title = doc.add_paragraph()
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("ITDMA Assignment 1 Working Pack").bold = True
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Risk Mitigation in CI/CD Pipelines for Safety-Critical Embedded Systems").italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Referencing style: Harvard author-date | Format target: Arial 12, 1.5 spacing | Due: 8 June 2026, 23:59")

    doc.add_heading("How to Use This Pack", level=1)
    for item in [
        "Use this document as a working pack and writing guide, not as a final submission.",
        "Write the final article in your own words, using the evidence prompts and Harvard citation cues.",
        "Replace every bracketed writing prompt before submission.",
        "Keep a live references list while drafting so in-text citations and reference entries match.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Research Focus", level=1)
    p = doc.add_paragraph()
    p.add_run("Topic: ").bold = True
    p.add_run("Risk Mitigation in CI/CD Pipelines for Safety-Critical Embedded Systems Using Automated Testing and Rollback Strategies.")
    p = doc.add_paragraph()
    p.add_run("Research question: ").bold = True
    p.add_run("How can automated testing and rollback strategies within CI/CD pipelines reduce deployment risks in safety-critical embedded systems?")
    doc.add_paragraph("Objectives:")
    for obj in [
        "Identify key risks associated with automated deployments in safety-critical systems.",
        "Examine automated testing techniques suitable for embedded CI/CD pipelines.",
        "Design a risk-aware CI/CD pipeline incorporating rollback and fail-safe mechanisms.",
    ]:
        doc.add_paragraph(obj, style="List Number")

    doc.add_heading("Distinction Argument", level=1)
    doc.add_paragraph(
        "The strongest argument is that safety-critical embedded CI/CD cannot simply copy web-style continuous deployment. "
        "It needs layered verification, traceable certification evidence, staged deployment, monitoring, and recovery mechanisms that select between rollback, fail-safe, fail-silent, fail-operational and graceful degradation modes."
    )

    doc.add_heading("Article Structure and Word Budget", level=1)
    rows = [
        ["Abstract", "150-200", "Write last. State the topic, method, core finding and recommendation."],
        ["Introduction", "350-450", "Problem, objectives, research question and roadmap. Use Zampetti et al. (2023), Dakkak et al. (2023), Baron and Louis (2023)."],
        ["Methodology", "450-600", "Databases, search strings, inclusion/exclusion criteria and matrix fields."],
        ["Literature Review", "650-800", "Theme-based synthesis of CI/CD, testing, certification and fail-safe recovery."],
        ["Results / Comparative Analysis", "800-1000", "Compare paper structures, methodologies, approaches, data collection, sampling, analysis and findings."],
        ["Discussion", "600-800", "Synthesis, gaps, implications and pipeline design response."],
        ["Conclusion and Venue", "350-500", "Answer research question, recommendations and ICST venue fit."],
    ]
    add_doc_table(doc, ["Section", "Words", "Writing instruction"], rows, [1.5, 0.8, 4.0])

    doc.add_heading("Common Research Paper Sections", level=1)
    rows = [[s["section"], s["purpose"], s["content"], s["style"], s["citation_hint"]] for s in SECTION_GUIDE]
    add_doc_table(doc, ["Section", "Purpose", "Typical content", "Style", "Citation cue"], rows, [1.1, 1.45, 1.55, 1.2, 1.2])

    doc.add_heading("Literature Matrix: Short Version", level=1)
    matrix_rows = []
    for paper in PAPERS:
        matrix_rows.append(
            [
                paper["id"],
                paper["citation"],
                paper["methodology"],
                paper["findings"],
                paper["gaps"],
                paper["use_in_article"],
            ]
        )
    add_doc_table(doc, ["ID", "Source", "Methodology", "Key finding", "Gap", "Use in article"], matrix_rows, [0.35, 1.0, 1.45, 1.65, 1.35, 1.5])

    doc.add_heading("Risk-Aware CI/CD Pipeline Model", level=1)
    pipeline_rows = []
    for index, step in enumerate(PIPELINE_MODEL, start=1):
        gate, description = step.split(": ", 1)
        pipeline_rows.append([f"Stage {index}", gate, description])
    add_doc_table(doc, ["Stage", "Gate", "Risk-control function"], pipeline_rows, [0.8, 1.6, 4.0])
    doc.add_paragraph(
        "Possible synthesis sentence: A risk-aware pipeline should treat rollback as one recovery option inside a broader assurance strategy, because safety-critical embedded systems may require degraded, fail-safe or fail-operational behaviour rather than a simple version reversal (Becker, Voss and Schätz, 2018)."
    )

    doc.add_heading("Harvard In-Text Citation Guide", level=1)
    for item in [
        "Narrative citation: Zampetti et al. (2023) found that CPS CI/CD requires balancing simulators and Hardware-in-the-Loop testing.",
        "Parenthetical citation: CPS pipelines require both software and hardware expertise (Zampetti et al., 2023).",
        "Two-author citation: Baker and Habli (2013) showed that mutation testing can reveal weaknesses missed by coverage criteria.",
        "Three-author citation in this school pack: use Becker, Voss and Schätz (2018) for the first and later mentions unless your lecturer requires 'et al.' after first mention.",
        "Multiple sources: Automated embedded testing should combine broad testing knowledge, virtual environments and physical validation (Garousi et al., 2018; Du et al., 2022; Barbie and Hasselbring, 2024).",
        "Avoid unsupported claims. If a sentence contains a factual claim from the literature, cite it.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Drafting Scaffold", level=1)
    scaffold_sections = [
        ("Abstract", [
            "[One sentence on the deployment-risk problem.]",
            "[One sentence on your literature-search method: eight peer-reviewed studies were analysed.]",
            "[One sentence on the main pattern: layered testing + staged deployment + recovery design.]",
            "[One sentence on your recommendation.]",
        ]),
        ("Introduction", [
            "Open with why faster deployment is attractive but risky in safety-critical embedded systems.",
            "Contrast ordinary CI/CD with CPS constraints using Zampetti et al. (2023).",
            "Introduce system-of-systems validation pressure using Dakkak et al. (2023).",
            "End with your research question and objectives.",
        ]),
        ("Methodology", [
            "State databases: Google Scholar, IEEE Xplore, ACM Digital Library, ScienceDirect and institutional repositories.",
            "State search terms: CI/CD, continuous deployment, cyber-physical systems, embedded software testing, safety-critical software, mutation testing, Hardware-in-the-Loop, digital twin prototype, rollback, fail-safe and graceful degradation.",
            "State inclusion criteria: peer-reviewed, directly relevant to CI/CD/testing/safety/recovery, preferably 2018-2024 except foundational safety-critical testing evidence.",
            "State analysis method: each paper was compared by section structure, method, approach, framework, data collection, sampling, analysis, findings, contradictions and gaps.",
        ]),
        ("Literature Review", [
            "Theme 1: CI/CD in CPS and systems-of-systems.",
            "Theme 2: automated embedded testing, mutation testing, HIL, simulation and digital twins.",
            "Theme 3: continuous certification and safety assurance evidence.",
            "Theme 4: rollback, fail-safe behaviour and graceful degradation.",
        ]),
        ("Results / Comparative Analysis", [
            "Compare article structures: empirical studies have explicit method/results sections; formal/design papers emphasise models and evaluation; reviews emphasise search/selection methods.",
            "Compare methodologies: interviews/surveys, case studies, engineering prototypes, systematic mapping, formal modelling and conceptual modelling.",
            "Group findings into confirming, extending and contradicting patterns.",
        ]),
        ("Discussion", [
            "Answer the research question directly: automated testing reduces deployment risk by catching defects earlier and producing evidence, while rollback/fail-safe mechanisms reduce harm after faults escape.",
            "Evaluate gaps: limited integrated research connecting CI/CD, certification artefacts, HIL/digital twins and rollback in one pipeline.",
            "Propose the six-stage risk-aware CI/CD model from this pack.",
        ]),
        ("Conclusion and Venue", [
            "Close with a concise answer and recommendations.",
            "Name ICST because it covers software testing, verification, validation, empirical studies and embedded software.",
        ]),
    ]
    for heading, bullets in scaffold_sections:
        doc.add_heading(heading, level=2)
        for bullet in bullets:
            doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("Rubric Checklist", level=1)
    for item in [
        "5-10 peer-reviewed sources included and all appear in the reference list.",
        "Common research paper sections are identified with purpose, content and writing style.",
        "Different article structures are compared instead of merely listed.",
        "Research methodologies, approaches, frameworks, data collection, sampling and analysis methods are identified.",
        "Findings are grouped into confirmations, extensions, contradictions, agreements and gaps.",
        "Recommendations connect directly to the chosen topic and research question.",
        "Harvard in-text citations match the final reference list.",
        "Final document uses Arial 12 and 1.5 spacing.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Harvard Reference List", level=1)
    for paper in sorted(PAPERS, key=lambda item: item["reference"].casefold()):
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.first_line_indent = Inches(-0.25)
        para.add_run(paper["reference"])

    doc.add_heading("Source Links", level=1)
    for paper in PAPERS:
        para = doc.add_paragraph()
        para.add_run(f"{paper['id']} {paper['citation']}: ")
        add_hyperlink(para, paper["url"], paper["url"])

    out = OUT_DIR / "ITDMA_Assignment_1_Working_Pack_Harvard.docx"
    doc.save(out)
    return out


def write_markdown() -> None:
    readme = """# ITDMA Assignment 1 Working Pack

Topic: **Risk Mitigation in CI/CD Pipelines for Safety-Critical Embedded Systems Using Automated Testing and Rollback Strategies**

Use this pack as a high-integrity writing workspace. It gives you the evidence base, Harvard citation cues, section plan and rubric checks. Your final submission should be written in your own words and revised by you.

## Files

- `ITDMA_Assignment_1_Working_Pack_Harvard.docx` - Word-ready scaffold and source pack.
- `literature_matrix_harvard.xlsx` - fillable source matrix.
- `literature_matrix_harvard.md` - readable source matrix.
- `article_scaffold_harvard.md` - section-by-section writing scaffold with in-text citation prompts.
- `rubric_distinction_checklist.md` - checklist against the assignment rubric.
- `harvard_reference_list.md` - Harvard references and citation examples.

## Core argument

Safety-critical embedded CI/CD should be risk-aware rather than speed-only. Automated tests reduce risk before deployment, but rollback and fail-safe mechanisms reduce harm when defects escape. The best answer combines simulation/digital twin testing, Hardware-in-the-Loop validation, mutation/coverage evidence, continuous certification artefacts, staged deployment, monitoring and recovery modes.
"""
    (OUT_DIR / "README.md").write_text(readme, encoding="utf-8")

    ref_lines = [
        "# Harvard Reference List and In-Text Citation Cues",
        "",
        "## In-text citation rules",
        "",
        "- Narrative: `Zampetti et al. (2023) argue that ...`",
        "- Parenthetical: `... in CPS CI/CD environments (Zampetti et al., 2023).`",
        "- Multiple sources: `... (Garousi et al., 2018; Du et al., 2022; Barbie and Hasselbring, 2024).`",
        "- Page numbers are only needed for direct quotations. Avoid direct quotes unless your lecturer expects them.",
        "",
        "## References",
        "",
    ]
    ref_lines.extend([f"- {paper['reference']}" for paper in sorted(PAPERS, key=lambda item: item["reference"].casefold())])
    (OUT_DIR / "harvard_reference_list.md").write_text("\n".join(ref_lines) + "\n", encoding="utf-8")

    matrix_lines = [
        "# Literature Matrix",
        "",
        "| ID | Source | Methodology | Approach | Data / sampling | Key findings | Gaps | How to use |",
        "|---|---|---|---|---|---|---|---|",
    ]
    for paper in PAPERS:
        matrix_lines.append(
            f"| {paper['id']} | {paper['citation']} {paper['in_text']} | {paper['methodology']} | {paper['approach']} | {paper['data_collection']} Sampling: {paper['sampling']} | {paper['findings']} | {paper['gaps']} | {paper['use_in_article']} |"
        )
    (OUT_DIR / "literature_matrix_harvard.md").write_text("\n".join(matrix_lines) + "\n", encoding="utf-8")

    scaffold = [
        "# Article Scaffold: Harvard Referencing",
        "",
        "Working title: **Risk Mitigation in CI/CD Pipelines for Safety-Critical Embedded Systems Using Automated Testing and Rollback Strategies**",
        "",
        "Research question: **How can automated testing and rollback strategies within CI/CD pipelines reduce deployment risks in safety-critical embedded systems?**",
        "",
        "Important: Replace every bracketed prompt with your own writing before submission.",
        "",
        "## Abstract (150-200 words)",
        "",
        "- [Problem context.]",
        "- [Method: eight peer-reviewed sources analysed.]",
        "- [Main synthesis.]",
        "- [Recommendation.]",
        "",
        "## Introduction (350-450 words)",
        "",
        "Use these citation moves:",
        "",
        "- CI/CD is useful but difficult in CPS because hardware, simulation and deployment constraints matter (Zampetti et al., 2023).",
        "- Continuous deployment in embedded systems-of-systems needs field validation, orchestration and monitoring (Dakkak et al., 2023).",
        "- Safety-critical contexts require continuous certification evidence, not only fast release cycles (Baron and Louis, 2023).",
        "",
        "Write in this order: context -> problem -> objectives -> research question -> article roadmap.",
        "",
        "## Methodology (450-600 words)",
        "",
        "Mention: Google Scholar, IEEE Xplore, ACM Digital Library, ScienceDirect and institutional repositories. Use search strings such as `CI/CD cyber-physical systems`, `continuous deployment embedded systems`, `safety-critical software mutation testing`, `Hardware-in-the-Loop continuous integration`, `digital twin prototypes embedded software`, and `graceful degradation automotive systems`.",
        "",
        "Inclusion criteria: peer-reviewed sources; relevance to CI/CD, embedded testing, safety-critical assurance or rollback/fail-safe design; enough methodological detail to compare research approach and findings.",
        "",
        "Exclusion criteria: vendor-only sources, non-peer-reviewed blog posts, sources without clear methodology, and sources about ordinary web CI/CD with no embedded or safety relevance.",
        "",
        "## Literature Review (650-800 words)",
        "",
        "Organise by themes, not by paper:",
        "",
        "1. CI/CD and deployment risks in cyber-physical and embedded systems: Zampetti et al. (2023), Dakkak et al. (2023).",
        "2. Automated embedded testing: Garousi et al. (2018), Du et al. (2022), Baker and Habli (2013), Barbie and Hasselbring (2024).",
        "3. Continuous certification: Baron and Louis (2023).",
        "4. Recovery and graceful degradation: Becker, Voss and Schätz (2018).",
        "",
        "## Results / Comparative Analysis (800-1000 words)",
        "",
        "Required comparisons:",
        "",
        "- Section structures: empirical studies vs reviews vs formal/design-science papers.",
        "- Methodologies: interviews/survey, case study, prototype evaluation, systematic mapping, formal modelling and conceptual formalisation.",
        "- Research approaches: interpretivist, positivist, pragmatic and design science.",
        "- Data collection, sampling and analysis methods.",
        "- Findings that confirm, extend and contradict earlier work.",
        "",
        "## Discussion (600-800 words)",
        "",
        "Answer the research question directly. A strong thesis sentence could be adapted in your own words:",
        "",
        "> Automated testing reduces deployment risk before release by increasing feedback quality and assurance evidence, while rollback and fail-safe strategies reduce post-release harm by controlling what happens when faults still escape.",
        "",
        "Discuss these gaps:",
        "",
        "- Few studies integrate CI/CD, safety certification, HIL/digital twins and rollback into one end-to-end model.",
        "- Rollback is under-theorised for safety-critical embedded systems because a simple version reversal may not be safe enough.",
        "- More empirical evidence is needed in regulated domains outside aerospace and telecommunications.",
        "",
        "## Conclusion and Conference Venue (350-500 words)",
        "",
        "Recommendation: submit to the **International Conference on Software Testing, Verification and Validation (ICST)** because it covers software testing, verification, validation, empirical studies, tools, embedded software and technology transfer.",
    ]
    (OUT_DIR / "article_scaffold_harvard.md").write_text("\n".join(scaffold) + "\n", encoding="utf-8")

    checklist = [
        "# Distinction Rubric Checklist",
        "",
        "## Literature Search and Source Selection (5%)",
        "- [ ] 8 peer-reviewed papers used.",
        "- [ ] Each source is relevant to CI/CD, embedded testing, safety-critical assurance or rollback/fail-safe design.",
        "",
        "## Identification of Research Paper Sections (10%)",
        "- [ ] Abstract, Introduction, Literature Review, Methodology, Results, Discussion, Conclusion and References are identified.",
        "- [ ] You compare how different paper types handle these sections.",
        "",
        "## Purpose, Content and Writing Style per Section (15%)",
        "- [ ] Each section has purpose, typical content and writing style.",
        "- [ ] Your explanation is clear and linked to examples from the selected papers.",
        "",
        "## Comparison of Article Structures (10%)",
        "- [ ] You compare empirical, review, formal, conceptual and design-science structures.",
        "- [ ] You explain why structure changes based on methodology.",
        "",
        "## Research Methodology Analysis (20%)",
        "- [ ] Methodology, approach, framework/theory, data collection, sampling and data analysis are identified for every paper.",
        "- [ ] You evaluate how methods help answer each paper's objective.",
        "",
        "## Findings Analysis and Literature Gaps (25%)",
        "- [ ] Findings are grouped into confirmation, extension, contradiction, agreement and gaps.",
        "- [ ] Discussion explains what the findings mean for a risk-aware CI/CD pipeline.",
        "- [ ] Gaps are named as methodological, contextual, technological, measurement or synthesis gaps.",
        "",
        "## Research Article Writing Quality (10%)",
        "- [ ] Article is 2000-4000 words, ideally 3000-3400.",
        "- [ ] Academic tone is clear, logical and not overcomplicated.",
        "",
        "## Referencing and Academic Integrity (5%)",
        "- [ ] Harvard in-text citations and reference list match.",
        "- [ ] All paraphrased ideas are cited.",
        "- [ ] Final text is written in your own words.",
    ]
    (OUT_DIR / "rubric_distinction_checklist.md").write_text("\n".join(checklist) + "\n", encoding="utf-8")


def create_xlsx() -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = "Literature Matrix"
    headers = [
        "ID",
        "Harvard citation",
        "In-text citation",
        "Peer-review status",
        "Purpose",
        "Article structure",
        "Methodology",
        "Approach",
        "Framework / theory",
        "Data collection",
        "Sampling",
        "Data analysis",
        "Key findings",
        "Confirms prior work",
        "Extends research",
        "Contradicts / qualifies",
        "Literature gap",
        "Use in your article",
        "URL",
        "Reference list entry",
    ]
    ws.append(headers)
    for paper in PAPERS:
        ws.append(
            [
                paper["id"],
                paper["citation"],
                paper["in_text"],
                paper["peer_review"],
                paper["purpose"],
                paper["structure"],
                paper["methodology"],
                paper["approach"],
                paper["framework"],
                paper["data_collection"],
                paper["sampling"],
                paper["analysis"],
                paper["findings"],
                paper["confirms"],
                paper["extends"],
                paper["contradicts"],
                paper["gaps"],
                paper["use_in_article"],
                paper["url"],
                paper["reference"],
            ]
        )
    header_fill = PatternFill("solid", fgColor="E8EEF5")
    for cell in ws[1]:
        cell.font = Font(bold=True, name="Arial", size=10)
        cell.fill = header_fill
        cell.alignment = Alignment(wrap_text=True, vertical="center", horizontal="center")
    widths = [8, 24, 20, 28, 42, 45, 42, 32, 42, 38, 34, 40, 48, 38, 38, 38, 42, 45, 40, 70]
    for idx, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(idx)].width = width
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.font = Font(name="Arial", size=10)
            cell.alignment = Alignment(wrap_text=True, vertical="top")
    ws.freeze_panes = "A2"
    out = OUT_DIR / "literature_matrix_harvard.xlsx"
    wb.save(out)
    return out


def main() -> None:
    write_markdown()
    create_xlsx()
    create_docx()


if __name__ == "__main__":
    main()
