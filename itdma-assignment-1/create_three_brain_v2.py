from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_assignment_pack import PAPERS, configure_styles


OUT_DIR = Path(__file__).resolve().parent
TITLE = "Risk Mitigation in CI/CD Pipelines for Safety-Critical Embedded Systems Using Automated Testing and Rollback Strategies"


ARTICLE = [
    (
        "Abstract",
        [
            "Safety-critical embedded systems increasingly rely on software updates, yet ordinary CI/CD practices cannot be transferred into these domains without additional assurance controls. A defective deployment may affect hardware behaviour, timing constraints, certification evidence and operational safety. This article reviews eight peer-reviewed studies on CI/CD in cyber-physical systems, continuous deployment in embedded systems-of-systems, automated embedded testing, mutation testing, digital twin prototypes, continuous certification and graceful degradation. The analysis argues that deployment risk should be managed through progressive assurance before deployment and controlled recovery after deployment. Automated testing reduces uncertainty by detecting defects earlier and producing repeatable evidence, while rollback, fail-safe and degraded-operation strategies reduce harm when defects still escape. The article proposes a risk-aware CI/CD pipeline that combines static and unit tests, simulation or digital twin testing, Hardware-in-the-Loop validation, mutation and coverage evidence, continuous certification artefacts, staged deployment, monitoring and predefined recovery rules.",
        ],
    ),
    (
        "Keywords",
        [
            "continuous integration; continuous deployment; safety-critical embedded systems; automated testing; rollback; graceful degradation; Hardware-in-the-Loop; digital twin prototypes",
        ],
    ),
    (
        "1. Introduction",
        [
            "Continuous integration and continuous deployment are normally valued because they shorten feedback cycles, expose defects earlier and support smaller releases. In ordinary software products, these practices can improve delivery speed and responsiveness. Safety-critical embedded systems create a different problem. Their software is connected to sensors, actuators, hardware timing, physical environments and regulatory evidence. A deployment failure may therefore produce more than inconvenience or downtime; it may create unsafe system behaviour or weaken the evidence needed to justify operational trust.",
            "The literature shows that cyber-physical and embedded systems need a more cautious CI/CD model than web-based software. Zampetti et al. (2023) found that CI/CD in cyber-physical systems depends on a balance between continuous and periodic builds, simulator use, Hardware-in-the-Loop (HIL) testing and combined hardware-software expertise. Dakkak et al. (2023) similarly show that continuous deployment in software-intensive systems-of-systems depends on field validation, orchestration, documentation and monitoring. Baron and Louis (2023) add that safety-critical avionics development requires certification evidence to be integrated into the development process rather than treated as a late-stage audit activity.",
            "This article investigates how automated testing and rollback strategies can reduce deployment risks in safety-critical embedded systems. The objectives are to identify key deployment risks, examine automated testing techniques suitable for embedded CI/CD pipelines, and propose a risk-aware pipeline that incorporates rollback and fail-safe mechanisms. The research question is: How can automated testing and rollback strategies within CI/CD pipelines reduce deployment risks in safety-critical embedded systems?",
        ],
    ),
    (
        "2. Methodology",
        [
            "This article uses a structured literature review approach. Sources were selected using structured search terms across academic databases and publisher repositories, including Google Scholar, IEEE Xplore, ACM Digital Library, ScienceDirect and institutional repositories. Search terms included continuous integration, continuous deployment, cyber-physical systems, embedded software testing, safety-critical software, Hardware-in-the-Loop, mutation testing, digital twin prototypes, continuous certification, rollback, fail-safe behaviour and graceful degradation. The review does not claim to be a full systematic literature review; instead, it uses a focused comparison matrix to evaluate a small but relevant set of peer-reviewed sources.",
            "The inclusion criteria were peer-reviewed status, direct relevance to CI/CD, embedded testing, safety-critical assurance or recovery behaviour, and enough methodological detail to compare research approach, data collection, sampling and analysis. Non-peer-reviewed vendor material, blog posts and general DevOps sources without embedded or safety-critical relevance were excluded. Baker and Habli (2013) was retained despite being older because it provides strong empirical evidence about mutation testing for safety-critical airborne software, which is directly relevant to test adequacy.",
            "Each source was analysed using the following fields: article structure, purpose, methodology, research approach, framework or model, data collection, sampling, data analysis, main findings, findings that confirm prior work, findings that extend existing research, findings that contradict or qualify earlier assumptions, literature gaps and relevance to the research question. This comparison method was influenced by the logic of systematic mapping, where literature is classified to identify patterns and gaps (Garousi et al., 2018).",
            "Eight sources were considered sufficient for this assignment because they cover the required breadth without turning the paper into a descriptive catalogue. The selected papers represent empirical practice evidence, industrial case-study evidence, systematic mapping, prototype evaluation, mutation-testing evaluation, conceptual modelling, certification-framework work and formal recovery modelling. This spread allows the article to compare not only what the studies found, but also how different research designs shape the strength and limits of their evidence.",
        ],
    ),
    (
        "3. Results: Article Structure and Methodology Comparison",
        [
            "The selected articles follow common research-paper sections, but their structures differ because they produce different kinds of evidence. Empirical software engineering papers such as Zampetti et al. (2023) and Dakkak et al. (2023) emphasise context, data collection, findings, discussion and validity. Systematic mapping work such as Garousi et al. (2018) gives more space to search strategy, inclusion criteria and classification. Design-science and framework papers such as Du et al. (2022), Barbie and Hasselbring (2024) and Baron and Louis (2023) foreground concepts, tools, models and demonstrations. Becker, Voss and Schätz (2018) uses a formal-methods structure because its evidence comes from modelling and constraint-based analysis.",
            "These structural differences matter for the assignment topic because no single methodology answers the research question alone. Interviews and case studies reveal real industrial barriers, but they are context-dependent. Prototype and design-science studies show feasibility, but they may not prove broad adoption. Formal modelling supports recovery logic, but it can simplify operational complexity. Systematic mapping provides breadth, but it is less specific about one pipeline architecture. A strong answer therefore needs synthesis across evidence types.",
            "The standard research-paper sections also serve different purposes. The abstract gives a compressed overview, the introduction frames the problem, the methodology explains how evidence was produced, the results section reports patterns, the discussion interprets meaning, the conclusion answers the research question, and the reference list makes the evidence traceable. The writing style changes across these sections: abstracts are concise, methods are objective and reproducible, while discussions are more evaluative and interpretive.",
            "The methodology comparison also indicates how each paper contributes to the three assignment objectives. CI/CD practice studies help identify deployment risks. Testing and mutation studies help evaluate automated testing techniques. Certification and graceful-degradation studies help design recovery and assurance mechanisms. The value of the literature is therefore cumulative: each source answers one part of the risk-mitigation problem, and the final pipeline model connects those parts.",
        ],
    ),
    (
        "4. Thematic Literature Review",
        [
            "The first theme is deployment risk in cyber-physical and embedded contexts. Zampetti et al. (2023) show that CPS CI/CD must manage simulator limitations, HIL availability, deployment complexity and the need for cross-disciplinary expertise. Dakkak et al. (2023) extend this into systems-of-systems by showing that continuous deployment changes field validation into an ongoing activity. Together, these studies show that deployment risk is not limited to whether a build passes. It also includes whether validation environments are realistic, whether monitoring is ready, and whether interacting systems can absorb change safely.",
            "The second theme is layered automated testing. Garousi et al. (2018) show that embedded software testing is broad and fragmented, which supports using multiple testing techniques rather than relying on one automated stage. Du et al. (2022) demonstrate this principle through a multitest CI platform for automotive systems using virtual and FPGA-supported verification. Barbie and Hasselbring (2024) add that digital twin prototypes can support repeatable virtual integration testing when physical hardware is scarce. The shared pattern is that faster virtual tests should be complemented by higher-fidelity hardware-aware tests.",
            "The third theme is test adequacy and assurance evidence. Baker and Habli (2013) found that mutation testing can reveal weaknesses in safety-critical airborne software even when traditional structural coverage has already been achieved. This finding is important because many CI pipelines treat a passing test suite as sufficient. In safety-critical embedded systems, the quality of the tests matters as much as the existence of tests. Baron and Louis (2023) extend the assurance argument by proposing continuous certification for safety-critical avionics, where traceability and certification artefacts are produced throughout development instead of being assembled after implementation.",
            "The fourth theme is recovery after faults escape testing. Rollback is useful, but it is not a complete safety strategy. Becker, Voss and Schätz (2018) show that fault-tolerant automotive systems may require fail-silent, fail-safe, fail-operational or degraded modes. This qualifies common CI/CD thinking because reverting to an older version may not be the safest response if the system is already in a hazardous physical state or if hardware configuration has changed. Recovery must therefore be designed as a decision process that uses monitoring evidence, operational context and predefined safety rules.",
            "Across the literature, a coherent pattern emerges: automated testing reduces risk before deployment, assurance artefacts make risk evidence auditable, and rollback or fail-safe mechanisms reduce harm after deployment. The contribution of this article is to connect these ideas into one risk-aware CI/CD pipeline for safety-critical embedded systems.",
        ],
    ),
    (
        "5. Findings, Gaps and Discussion",
        [
            "The findings that confirm prior research are mainly about the importance of early and repeated verification. Zampetti et al. (2023), Garousi et al. (2018) and Du et al. (2022) all support the idea that automated testing improves feedback and helps detect defects earlier. However, they also show that embedded systems require hardware-aware testing. This confirms the general value of CI/CD while qualifying it for safety-critical contexts.",
            "The findings that extend existing research move beyond simple build automation. Dakkak et al. (2023) extend CI/CD thinking by showing that continuous deployment in systems-of-systems requires field validation and orchestration. Baron and Louis (2023) extend Agile and CI practice into certification by proposing continuous certification evidence. Barbie and Hasselbring (2024) extend simulation-based testing by formalising digital twin prototypes, while Becker, Voss and Schätz (2018) extend rollback thinking into graceful degradation and redundancy-aware recovery.",
            "The main contradiction is not that the papers disagree about the value of CI/CD, but that they challenge the assumption that mainstream CI/CD can be transferred directly into safety-critical embedded systems. In web systems, rollback may be treated as a normal operational safeguard. In embedded safety-critical systems, rollback is only one possible response. Depending on the fault, a safer response may be degraded operation, component isolation, fail-safe shutdown or fail-operational redundancy. This contradiction exists because the consequence model is different: embedded systems interact with physical environments, and safety depends on system state as well as software version.",
            "The reviewed papers also share an important area of agreement: pure web-style CI/CD is insufficient for safety-critical embedded systems. The articles differ in domain and method, but they agree that embedded deployment requires extra attention to hardware realism, validation evidence and system-level behaviour. This agreement directly influences the article's argument because it shifts CI/CD from a speed-focused delivery practice to a controlled assurance process.",
            "Several literature gaps are visible when the assignment appendix categories are applied. The methodological gap is that many studies focus on a single technique or domain rather than evaluating complete end-to-end pipelines. The theoretical or conceptual gap is that CI/CD, certification and graceful degradation are rarely integrated into one shared model. The contextual gap is that evidence comes mainly from cyber-physical systems, telecommunications, automotive and avionics, so transfer to medical, industrial or other regulated domains requires caution. The data or technology gap is that digital twin prototypes and HIL environments are promising, but their fidelity and certification acceptance remain open questions. The measurement or evaluation gap is that the literature lacks a standard metric for deployment readiness that combines test adequacy, hardware realism, certification evidence, monitoring and recovery readiness. The temporal gap is visible in the reliance on Baker and Habli's older but still relevant mutation-testing evidence, which needs updating for modern high-velocity CI/CD tooling. The outcome or impact gap is that many studies focus on technical feasibility without long-term evidence of field safety outcomes. The synthesis gap is the most important: CI/CD, testing, certification and fault-tolerance research exist, but they are rarely integrated into one risk-aware deployment model.",
            "Conceptually, the findings refine ordinary CI/CD from a delivery automation framework into a sociotechnical assurance process. The pipeline is not only a build mechanism; it is also a way of coordinating people, tools, evidence, hardware environments and recovery decisions. This theoretical implication is visible across the interview evidence in Zampetti et al. (2023), the certification framework in Baron and Louis (2023), and the formal recovery logic in Becker, Voss and Schätz (2018).",
            "The practical implication is that risk mitigation should be evaluated before, during and after deployment. Before deployment, the question is whether test evidence is strong enough and realistic enough for the change being released. During deployment, the question is whether staged rollout and monitoring can detect abnormal behaviour early. After deployment, the question is whether the system can move to an appropriate recovery state quickly enough to prevent harm. This three-part evaluation makes the discussion more specific than a general recommendation to add more testing.",
            "These findings influence the proposed research by shifting the focus from speed to progressive assurance. A safety-critical embedded pipeline should not ask only whether code can be deployed automatically. It should ask what evidence is required before deployment, what monitoring is needed during rollout, and what recovery decision should occur if the deployment behaves unexpectedly. This directly addresses the research question because automated testing reduces uncertainty before release, while rollback and fail-safe mechanisms limit harm after release.",
        ],
    ),
    (
        "6. Proposed Risk-Aware CI/CD Pipeline",
        [
            "A risk-aware CI/CD pipeline should begin with prevention gates. The first gate should include static analysis, build checks, unit tests and requirements traceability checks. These controls are fast and should run frequently. The second gate should include simulation and digital twin prototype tests because they provide repeatable integration feedback before scarce physical hardware is used (Barbie and Hasselbring, 2024). The third gate should include HIL or FPGA-supported validation where timing, physical interaction or hardware behaviour creates risk (Zampetti et al., 2023; Du et al., 2022).",
            "The next layer should create assurance evidence. Mutation testing, structural coverage review, configuration evidence and certification artefact generation should be part of the pipeline rather than late documentation work. Baker and Habli (2013) show why test adequacy matters, while Baron and Louis (2023) show why certification evidence should be continuous. The deployment layer should then use staged rollout, monitoring and field validation, especially in systems-of-systems where changes may affect interactions between components (Dakkak et al., 2023).",
            "The final layer should define controlled recovery. Rollback is suitable when monitoring detects a software regression, the previous version is known to be safe, and the hardware or configuration state remains compatible. Feature disablement is more suitable when a non-critical function fails but the core safety function remains stable. Component isolation or degraded mode is appropriate when continued operation is possible only with reduced capability. Fail-safe shutdown is needed when continued operation would be unsafe, while fail-operational redundancy is needed when stopping would create a greater hazard and redundant capacity exists. This decision logic is based on the graceful-degradation reasoning of Becker, Voss and Schätz (2018), but it is proposed here as a CI/CD recovery model that still requires empirical validation.",
            "Monitoring is the link between deployment and recovery. For a safety-critical embedded release, monitoring should include software health signals, hardware interaction signals, timing deviations, failed safety checks, configuration drift and field-validation outcomes. These signals should not merely create alerts for developers; they should trigger predefined decisions about rollout pause, rollback, feature disablement or safe-state transition. Without this connection, monitoring becomes passive observation rather than risk mitigation.",
            "In this model, automated testing and rollback work together rather than separately. Testing reduces the chance of failure before deployment, certification artefacts make assurance evidence auditable, staged monitoring limits exposure during release, and recovery rules reduce the consequences of failure after deployment. The model is therefore best understood as a proposed synthesis of the reviewed literature, not as a fully validated industrial standard.",
        ],
    ),
    (
        "7. Journal or Conference Venue",
        [
            "A suitable venue for this article would be the International Conference on Software Testing, Verification and Validation (ICST). ICST is appropriate because its scope includes software testing, verification, validation, empirical studies, tools, embedded software and technology transfer (ICST Steering Committee, 2023). The article's focus on automated testing, HIL, mutation testing, digital twin test environments and deployment risk in embedded systems fits this venue better than a broad software engineering venue. ICST should be treated as the venue source only, not as one of the eight peer-reviewed research papers.",
        ],
    ),
    (
        "8. Conclusion",
        [
            "Automated testing and rollback strategies can reduce deployment risks in safety-critical embedded systems when they are implemented as part of a broader risk-aware CI/CD pipeline. The literature shows that embedded CI/CD requires more than fast builds and frequent releases. It requires layered automated testing, hardware-aware validation, test adequacy evidence, continuous certification artefacts, staged deployment, monitoring and predefined recovery behaviour. Rollback is valuable, but it should be treated as one recovery option rather than a universal solution. In safety-critical systems, degraded operation, fail-safe shutdown or fail-operational redundancy may be safer than simple version reversal. The main contribution of this article is therefore a synthesis model: progressive assurance before deployment plus controlled recovery after deployment.",
        ],
    ),
]


METHODOLOGY_ROWS = [
    ["Zampetti et al. (2023)", "Interviews and survey validation", "Interpretivist / pragmatic", "CPS CI/CD practice barriers", "Interviews, card sorting, member checking and survey", "Shows CPS-specific deployment and HIL/simulator risks; context-dependent."],
    ["Dakkak et al. (2023)", "Industrial case study", "Pragmatic", "Continuous deployment in systems-of-systems", "Case evidence from Ericsson 3G RAN", "Shows field validation and orchestration needs; telecom-specific."],
    ["Du et al. (2022)", "Engineering prototype", "Design science", "Automated multitest CI platform", "Automotive platform demonstration", "Shows feasibility of virtual/FPGA CI testing; not broad industry proof."],
    ["Baker and Habli (2013)", "Empirical evaluation", "Positivist", "Mutation testing and test adequacy", "Safety-critical C and Ada software", "Strengthens assurance argument; older and not CI/CD-specific."],
    ["Garousi et al. (2018)", "Systematic mapping", "Positivist synthesis", "Embedded testing taxonomy", "312 selected technical papers", "Justifies layered testing; broad rather than pipeline-specific."],
    ["Barbie and Hasselbring (2024)", "Conceptual modelling and applications", "Design science", "Digital twin prototypes", "Field/lab application evidence", "Supports virtual integration tests; fidelity/certification remain open."],
    ["Baron and Louis (2023)", "Framework and tooling proposal", "Design science / pragmatic", "Continuous certification", "Industrial avionics case", "Supports continuous artefacts; avionics-specific."],
    ["Becker, Voss and Schätz (2018)", "Formal modelling", "Formal analytical", "Graceful degradation and redundancy", "Constructed automotive scenario", "Supports fail-safe recovery logic; not direct rollback evidence."],
]


SECTION_ROWS = [
    ["Abstract", "Summarises the whole article", "Problem, method, main synthesis and recommendation", "Concise and objective", "All papers use it to compress contribution."],
    ["Introduction", "Frames the problem and aim", "Context, gap, objectives and research question", "Persuasive but evidence-based", "Empirical papers stress industrial context; formal papers stress modelling problem."],
    ["Literature / Background", "Positions the study in prior work", "Definitions, related work and theoretical or technical concepts", "Analytical and source-based", "Review papers make this broader; design papers make it tool/model focused."],
    ["Methodology", "Explains how evidence was produced", "Research design, data collection, sampling and analysis", "Objective and reproducible", "Systematic mapping is most explicit; formal/design papers explain models or prototypes."],
    ["Results", "Reports findings or model outputs", "Themes, classifications, measurements, cases or formal outcomes", "Evidence-led and organised", "Empirical papers report participant/case findings; formal papers report model implications."],
    ["Discussion", "Interprets meaning", "Agreement, contradictions, implications, limitations and gaps", "Evaluative and comparative", "Strongest place to connect findings to the research question."],
    ["Conclusion / References", "Closes and documents evidence", "Answer to research question, future work and full sources", "Clear, restrained and accurate", "References make Harvard in-text claims traceable."],
]


FINDING_ROWS = [
    ["Confirm", "Layered automated testing improves feedback but must account for embedded constraints.", "Zampetti et al. (2023); Garousi et al. (2018); Du et al. (2022)", "Automated testing is useful only when paired with hardware-aware validation."],
    ["Extend", "Continuous deployment expands into field validation, certification evidence and digital twin testing.", "Dakkak et al. (2023); Baron and Louis (2023); Barbie and Hasselbring (2024)", "The pipeline must produce both test results and assurance artefacts."],
    ["Contradict / qualify", "Rollback is not always the safest response in physical systems.", "Becker, Voss and Schätz (2018)", "Recovery should include fail-safe, fail-operational and degraded modes."],
    ["Agreement", "Pure web-style CI/CD is insufficient for safety-critical embedded systems.", "Across the reviewed sources", "CI/CD must become a controlled assurance and recovery process."],
    ["Gap", "Few studies integrate CI/CD, HIL/digital twins, mutation testing, certification and recovery in one model.", "Across reviewed literature", "This article's contribution is a synthesis model for risk-aware CI/CD."],
]


PIPELINE_ROWS = [
    ["1", "Commit and build", "Static analysis, build checks, unit tests and traceability checks", "Fast prevention"],
    ["2", "Simulation / digital twin", "Repeatable virtual integration tests", "Early integration confidence"],
    ["3", "HIL / FPGA validation", "Hardware-aware timing and physical-interaction tests", "Environmental realism"],
    ["4", "Assurance evidence", "Mutation, coverage, configuration and certification artefacts", "Auditable safety evidence"],
    ["5", "Staged deployment", "Limited rollout, monitoring and field validation", "Controlled exposure"],
    ["6", "Recovery decision", "Rollback, feature disablement, isolation, degraded mode or fail-safe behaviour", "Post-deployment harm reduction"],
]


EXTRA_REFERENCES = [
    "ICST Steering Committee (2023) 'The International Conference on Software Testing, Verification and Validation - ICST'. Available at: https://icstconference.github.io/ (Accessed: 4 June 2026).",
]


AI_DECLARATION = """I used OpenAI Codex as the primary workspace assistant to plan the assignment, inspect the assignment brief, structure the literature matrix, generate and revise a model draft, create Harvard-style reference support, and verify the Word document layout. I also used the three-brain workflow for advisory critique: Claude CLI was used in read-only mode for adversarial review, Gemini CLI was used in read-only mode for broad-context/rubric review, and Codex remained responsible for final synthesis, document generation, and verification. I reviewed, edited and rewrote the final submission in my own words, verified the sources, and take responsibility for the submitted content."""


def add_page_number_footer(doc: Document) -> None:
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("ITDMA3-22 Assignment 1 | Page ")
    run.font.size = Pt(9)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D0D7DE")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade_cell(cell, "E8EEF5")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(8)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cell = cells[index]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(value)
            run.font.name = "Arial"
            run.font.size = Pt(7.2)
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Inches(width)
    doc.add_paragraph()


def add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Inches(0.25)
    paragraph.add_run(text)


def article_word_count() -> int:
    body = " ".join(paragraph for _, paragraphs in ARTICLE for paragraph in paragraphs)
    tables = " ".join(" ".join(row) for row in METHODOLOGY_ROWS + FINDING_ROWS + PIPELINE_ROWS)
    return len(re.findall(r"\b[\w'-]+\b", body + " " + tables))


def sorted_harvard_references() -> list[str]:
    entries = [paper["reference"] for paper in PAPERS] + EXTRA_REFERENCES
    return sorted(entries, key=lambda entry: entry.casefold())


def create_docx() -> Path:
    doc = Document()
    configure_styles(doc)
    add_page_number_footer(doc)

    title = doc.add_paragraph()
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(TITLE).bold = True

    for item in [
        "Student name: [Insert your name]",
        "Student number: [Insert your student number]",
        "Campus: [Insert campus name]",
        "Module: ITDMA3-22 Research Design and Methodology",
        "Referencing style: Harvard author-date",
        f"Approximate article word count excluding references: {article_word_count()}",
    ]:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(item)

    for heading, paragraphs in ARTICLE:
        doc.add_heading(heading, level=1)
        for paragraph in paragraphs:
            add_body_paragraph(doc, paragraph)
        if heading == "3. Results: Article Structure and Methodology Comparison":
            doc.add_heading("Table 1: Research Section Purpose, Content and Style", level=2)
            add_table(
                doc,
                ["Section", "Purpose", "Typical content", "Writing style", "Comparison note"],
                SECTION_ROWS,
                [1.0, 1.35, 1.5, 1.1, 1.55],
            )
            doc.add_heading("Table 2: Methodology and Evidence-Strength Matrix", level=2)
            add_table(
                doc,
                ["Source", "Methodology", "Approach", "Framework / model", "Data / sampling", "Contribution and limitation"],
                METHODOLOGY_ROWS,
                [1.0, 1.1, 0.95, 1.2, 1.25, 1.85],
            )
        if heading == "5. Findings, Gaps and Discussion":
            doc.add_heading("Table 3: Findings and Gap Synthesis", level=2)
            add_table(
                doc,
                ["Category", "Synthesis", "Evidence", "Influence on research"],
                FINDING_ROWS,
                [0.9, 2.0, 1.7, 1.9],
            )
        if heading == "6. Proposed Risk-Aware CI/CD Pipeline":
            doc.add_heading("Table 4: Proposed Pipeline Gates", level=2)
            add_table(
                doc,
                ["Stage", "Gate", "Main controls", "Risk function"],
                PIPELINE_ROWS,
                [0.55, 1.45, 2.75, 1.75],
            )

    doc.add_heading("References", level=1)
    for reference in sorted_harvard_references():
        ref = doc.add_paragraph()
        ref.paragraph_format.left_indent = Inches(0.25)
        ref.paragraph_format.first_line_indent = Inches(-0.25)
        ref.add_run(reference)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Arial"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
            if run.font.size is None:
                run.font.size = Pt(12)

    out = OUT_DIR / "ITDMA3-22_Assignment_1_Three_Brain_V2_Harvard.docx"
    doc.save(out)
    return out


def create_declaration() -> Path:
    out = OUT_DIR / "ai_declaration_three_brain_v2.md"
    out.write_text(AI_DECLARATION + "\n", encoding="utf-8")
    return out


def create_rubric_checklist() -> Path:
    checklist = """# Three-Brain V2 Rubric Checklist

- Literature search: 8 peer-reviewed sources used; ICST is venue only.
- Sections: abstract, introduction, methodology, article-structure/method comparison, literature review, findings/gaps discussion, pipeline proposal, venue, conclusion, references.
- Methodology: Table 1 explicitly identifies methodology, approach, model/framework, data/sampling, contribution and limitation for every source.
- Findings/gaps: Table 2 and Section 5 explicitly cover confirm, extend, contradict/qualify and gap categories.
- Gap labels: methodological, contextual, technology, measurement and synthesis gaps are named in prose.
- Recommendations: Section 6 turns literature findings into a concrete risk-aware CI/CD pipeline.
- Harvard: author-date in-text citations are used and the reference list matches.
- Final student pass: fill student details, verify source links, rewrite in your own voice, submit AI declaration separately.
"""
    out = OUT_DIR / "three_brain_v2_rubric_checklist.md"
    out.write_text(checklist, encoding="utf-8")
    return out


if __name__ == "__main__":
    print(create_docx())
    print(create_declaration())
    print(create_rubric_checklist())
