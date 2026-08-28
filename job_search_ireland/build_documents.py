from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "outputs"
OUTPUT.mkdir(parents=True, exist_ok=True)

NAVY = RGBColor(20, 36, 55)
TEAL = RGBColor(0, 117, 117)
SLATE = RGBColor(66, 76, 89)
LIGHT = "E8F2F2"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.add_run(" | Page ")
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def configure_document(doc, compact=False):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.58 if compact else 0.72)
    section.bottom_margin = Inches(0.55 if compact else 0.72)
    section.left_margin = Inches(0.68 if compact else 0.82)
    section.right_margin = Inches(0.68 if compact else 0.82)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.2 if compact else 10.2)
    normal.font.color.rgb = SLATE
    normal.paragraph_format.space_after = Pt(2.5 if compact else 5)
    normal.paragraph_format.line_spacing = 1.03 if compact else 1.08

    for name, size, before, after in [
        ("Heading 1", 11.2 if compact else 15, 8, 3),
        ("Heading 2", 10.2 if compact else 12, 6, 2),
        ("Heading 3", 9.5 if compact else 10.5, 4, 1),
    ]:
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = TEAL
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "CV Entry" not in doc.styles:
        style = doc.styles.add_style("CV Entry", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Aptos"
        style.font.size = Pt(9.4)
        style.font.bold = True
        style.font.color.rgb = NAVY
        style.paragraph_format.space_before = Pt(3)
        style.paragraph_format.space_after = Pt(1)
        style.paragraph_format.keep_with_next = True

    bullet = doc.styles["List Bullet"]
    bullet.font.name = "Aptos"
    bullet.font.size = Pt(9.0 if compact else 10)
    bullet.font.color.rgb = SLATE
    bullet.paragraph_format.left_indent = Inches(0.18)
    bullet.paragraph_format.first_line_indent = Inches(-0.12)
    bullet.paragraph_format.space_after = Pt(1.3 if compact else 3)
    bullet.paragraph_format.line_spacing = 1.0 if compact else 1.06

    footer = section.footer.paragraphs[0]
    footer.text = "Joshua Nehohwa | Ireland job-search materials"
    footer.runs[0].font.name = "Aptos"
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = SLATE
    add_page_number(footer)


def add_bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def add_labelled_paragraph(doc, label, text):
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.bold = True
    r.font.color.rgb = NAVY
    p.add_run(text)
    return p


def build_cv():
    doc = Document()
    configure_document(doc, compact=True)
    doc.sections[0].footer.paragraphs[0].text = ""

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("JOSHUA NEHOHWA")
    r.font.name = "Aptos Display"
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(1)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("GRADUATE FULL-STACK SOFTWARE DEVELOPER")
    r.font.name = "Aptos"
    r.font.size = Pt(10.2)
    r.font.bold = True
    r.font.color.rgb = TEAL
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.add_run("Cape Town, South Africa | Open to relocate to Ireland | ")
    p.add_run("nehohwajoshua@gmail.com | +27 63 857 3965\n")
    p.add_run("linkedin.com/in/joshua-nehohwa-b4b97b229 | github.com/jnehohwa")

    doc.add_heading("PROFILE", level=1)
    doc.add_paragraph(
        "Software Engineering student and junior full-stack developer with hands-on experience building web platforms, "
        "data-backed workflows, and operational automations using React, Next.js, TypeScript, PHP/MySQL, Supabase, "
        "WordPress, and AWS services. Achieved an 86% academic average, led a university developer community, delivered "
        "a public website and privacy-conscious intake workflow during a recent internship, and currently tutors first-year "
        "IT students. Expected to complete all degree requirements on 10 November 2026, available to relocate from "
        "April 2027, and open to graduate or junior software roles across Ireland."
    )

    doc.add_heading("TECHNICAL SKILLS", level=1)
    add_labelled_paragraph(doc, "Core development: ", "TypeScript, JavaScript, React, Next.js, Node.js, HTML, CSS, Python, PHP, Java")
    add_labelled_paragraph(doc, "Data and platforms: ", "PostgreSQL, MySQL, Supabase, Clerk Auth, WordPress, Elementor, Airtable, Jotform")
    add_labelled_paragraph(doc, "Cloud and delivery: ", "AWS EC2, S3, IAM and Lambda; Git, Linux, Docker fundamentals, Make.com, Google Apps Script")

    doc.add_heading("EDUCATION", level=1)
    p = doc.add_paragraph(style="CV Entry")
    p.add_run("BSc Information Technology (Software Engineering) | Eduvos, Cape Town").bold = True
    p.add_run(" | Degree requirements expected complete 10 Nov 2026")
    doc.add_paragraph("Academic average: 86% | Graduation ceremony expected May 2027 | Relevant study: data structures and algorithms, software architecture, mobile development, network security, cloud technologies, mathematics, and AI ethics.")

    doc.add_heading("EXPERIENCE", level=1)
    p = doc.add_paragraph(style="CV Entry")
    p.add_run("Digital Systems and Web Development Intern | Nature's Valley Trust / KuCoNa").bold = True
    p.add_run(" | Jun-Aug 2026")
    add_bullet(doc, "Built and launched a responsive WordPress/Elementor website spanning three programme hubs, 15+ public pages, registration pathways, and a 14-partner directory.")
    add_bullet(doc, "Designed a privacy-conscious workflow linking WordPress, Jotform, Make.com, and Airtable for youth registration, integration logging, attendance support, and duplicate prevention.")
    add_bullet(doc, "Cleaned 1,037 demonstration records while preserving the production Airtable schema and integrations.")
    add_bullet(doc, "Developed a 47-question bilingual community survey with Google Forms, Apps Script, and linked Sheets analytics, then produced handover documentation and staff training resources.")
    add_bullet(doc, "Live website: https://kucona.org.za")

    p = doc.add_paragraph(style="CV Entry")
    p.add_run("ITMTA1-B22 Tutor | Eduvos").bold = True
    p.add_run(" | May 2026-Present")
    add_bullet(doc, "Lead weekly two-hour Microsoft Teams tutorials, support exam logistics, coordinate availability with staff, and track allocated tutoring hours for first-year IT students.")

    doc.add_page_break()
    p = doc.add_paragraph(style="CV Entry")
    p.add_run("Freelance AI Data Specialist | Data Annotations AI").bold = True
    p.add_run(" | Remote | Jan-Nov 2024")
    add_bullet(doc, "Improved large-language-model training data by transforming raw HTML and text into high-quality annotations and evaluating responses against complex, evolving guidelines.")
    add_bullet(doc, "Documented edge cases and actionable quality feedback while meeting independent quality and throughput requirements on Google and Microsoft contract work.")

    p = doc.add_paragraph(style="CV Entry")
    p.add_run("Seasonal Lift Attendant | Snowshoe Mountain").bold = True
    p.add_run(" | West Virginia, USA")
    add_bullet(doc, "Worked 3 Dec 2024-28 Feb 2025 and 4 Dec 2025-1 Mar 2026; contracted to return 3 Dec 2026-1 Mar 2027.")
    add_bullet(doc, "Managed guest safety and lift operations in high-pressure winter conditions and received multiple top-performer recognitions for reliability and service.")

    doc.add_heading("SELECTED PROJECTS AND LEADERSHIP", level=1)
    p = doc.add_paragraph(style="CV Entry")
    p.add_run("Innosimm ERP System | Developer").bold = True
    p.add_run(" | 2025-Present")
    add_bullet(doc, "Designing a modular Next.js, Supabase, and Clerk system for inventory, sales, purchasing, quotations, invoices, and finance tracking for a vehicle and tyre business.")
    add_bullet(doc, "Modeling stock movement and document workflows to replace fragmented spreadsheet and WhatsApp processes with a maintainable data-backed system.")
    add_bullet(doc, "Status: in development; no public deployment yet.")

    p = doc.add_paragraph(style="CV Entry")
    p.add_run("HackJam Innovation Platform | Team Leader").bold = True
    p.add_run(" | 2025")
    add_bullet(doc, "Led a team building a React, Next.js, TypeScript, and Tailwind platform for idea submission, voting, mentor feedback, and gamification; placed 4th university-wide.")
    add_bullet(doc, "Coordinated user research, prototyping, technical prioritisation, and the final product pitch across the project team.")
    add_bullet(doc, "Repository: https://github.com/jnehohwa/HackJam-Shark-Tank")

    p = doc.add_paragraph(style="CV Entry")
    p.add_run("KasiSwap Marketplace | Developer").bold = True
    p.add_run(" | 2026")
    add_bullet(doc, "Built React/TypeScript and PHP/MySQL marketplace flows covering listings, authentication-aware actions, order states, messaging, disputes, reviews, and admin moderation.")
    add_bullet(doc, "Produced deployment notes, code evidence, and structured technical documentation for assessment and maintainability.")
    add_bullet(doc, "Demo: https://kasiswap.free.nf/?i=1")

    p = doc.add_paragraph(style="CV Entry")
    p.add_run("Chairperson | Vossie DevClub").bold = True
    p.add_run(" | 2025-Present")
    add_bullet(doc, "Lead technical workshops, student developer activities, mentorship initiatives, hackathons, and programming outreach for high-school learners.")
    add_bullet(doc, "Coordinate student-focused activities and communicate technical concepts to audiences with different levels of experience.")

    doc.add_heading("CERTIFICATIONS AND ACHIEVEMENTS", level=1)
    add_bullet(doc, "AWS Certified Solutions Architect - Associate | https://www.credly.com/badges/43660ab1-f99a-4bbe-ab14-261dfb57cfcd")
    add_bullet(doc, "AWS Certified Cloud Practitioner | https://www.credly.com/badges/61737c5d-678d-4633-80d4-7ea1c5d24b80")
    add_bullet(doc, "Golden Key International Honour Society | Top 15% of university cohort, 2025")
    add_bullet(doc, "IELTS Academic: 8.5 overall / CEFR C2 (Listening 8.5, Reading 9.0, Writing 6.5, Speaking 9.0)")

    path = OUTPUT / "Joshua_Nehohwa_Ireland_Graduate_CV.docx"
    doc.save(path)
    return path


def build_application_pack():
    doc = Document()
    configure_document(doc, compact=False)

    p = doc.add_paragraph()
    r = p.add_run("IRELAND SOFTWARE JOB APPLICATION KIT")
    r.font.name = "Aptos Display"
    r.font.size = Pt(21)
    r.font.bold = True
    r.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(2)
    p = doc.add_paragraph("Joshua Nehohwa | Graduate and junior full-stack roles | August 2026")
    p.runs[0].font.color.rgb = TEAL
    p.runs[0].font.bold = True

    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    widths = [Inches(1.55), Inches(5.1)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]
    table.cell(0, 0).text = "Target"
    table.cell(0, 1).text = "Graduate and junior full-stack software-development roles across Ireland"
    table.cell(1, 0).text = "Workflow"
    table.cell(1, 1).text = "Research and tailor -> Joshua reviews -> submit -> track and follow up"
    for row in table.rows:
        set_cell_shading(row.cells[0], LIGHT)
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_heading("1. Confirmed candidate facts", level=1)
    for item in [
        "Degree requirements and credits expected complete on 10 November 2026; graduation ceremony expected in May 2027.",
        "Earliest relocation and Ireland start availability: April 2027.",
        "No current Irish or EU work rights; the role must support the appropriate employment-permit route.",
        "Passport valid until August 2032.",
        "Snowshoe Mountain: 3 Dec 2024-28 Feb 2025; 4 Dec 2025-1 Mar 2026; contracted return 3 Dec 2026-1 Mar 2027.",
        "KuCoNa: https://kucona.org.za | HackJam: https://github.com/jnehohwa/HackJam-Shark-Tank",
        "KasiSwap: https://kasiswap.free.nf/?i=1 | Innosimm: in development with no public deployment.",
        "AWS Solutions Architect - Associate: https://www.credly.com/badges/43660ab1-f99a-4bbe-ab14-261dfb57cfcd",
        "AWS Cloud Practitioner: https://www.credly.com/badges/61737c5d-678d-4633-80d4-7ea1c5d24b80",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Still to verify before submission", level=2)
    for item in [
        "Official completion letter or final transcript availability after 10 November 2026, because the ceremony and certificate may follow later.",
        "Official institution and degree wording exactly as printed on academic records.",
        "KasiSwap availability from an employer-facing browser; the supplied URL did not respond during the 20 August 2026 automated check.",
        "AWS badge issue dates, driving-licence details, relocation funding, and restrictions on publishing internship screenshots or code.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("2. Positioning and evidence bank", level=1)
    doc.add_heading("Core proposition", level=2)
    doc.add_paragraph(
        "A high-performing Software Engineering student who has already delivered real web, data, and automation outcomes; "
        "can work across frontend, backend-connected workflows, documentation, and stakeholder handover; and brings AWS "
        "foundations, international work experience, tutoring, and technical leadership."
    )
    evidence = [
        ("KuCoNa website", "15+ public pages, three programme hubs, 14-partner directory, registration pathways."),
        ("Operational workflow", "Connected WordPress, Jotform, Make.com, and Airtable with privacy and duplicate prevention in mind."),
        ("Data quality", "Removed 1,037 demonstration records without damaging the live schema or integrations."),
        ("Survey delivery", "Built a 47-question bilingual survey with Apps Script and linked analytics."),
        ("HackJam", "Led product design and frontend delivery; team placed 4th university-wide."),
        ("Teaching and leadership", "Runs weekly tutorials and leads developer workshops, outreach, mentoring, and hackathons."),
    ]
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.autofit = False
    t.columns[0].width = Inches(1.55)
    t.columns[1].width = Inches(5.1)
    t.rows[0].cells[0].text = "Evidence"
    t.rows[0].cells[1].text = "Defensible proof point"
    set_repeat_table_header(t.rows[0])
    for c in t.rows[0].cells:
        set_cell_shading(c, "007575")
        for run in c.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    for label, value in evidence:
        row = t.add_row().cells
        row[0].text = label
        row[1].text = value

    doc.add_heading("3. Reusable LinkedIn copy", level=1)
    doc.add_heading("Headline", level=2)
    doc.add_paragraph("Graduate Full-Stack Developer | React, Next.js, TypeScript, Supabase and AWS | Software Engineering Student | Open to Relocate to Ireland")
    doc.add_heading("About", level=2)
    doc.add_paragraph(
        "I am a Software Engineering student expecting to complete all degree requirements on 10 November 2026, with an 86% academic average and practical "
        "experience building web platforms, workflow automations, and data-backed operational tools. During my recent "
        "internship with Nature's Valley Trust / KuCoNa, I helped launch a multi-page public website, connected digital "
        "registration systems, cleaned more than 1,000 demonstration records without disrupting the underlying schema, "
        "and created a bilingual survey with linked analytics.\n\n"
        "My core stack includes React, Next.js, TypeScript, JavaScript, PHP/MySQL, Supabase, and AWS fundamentals. I also "
        "lead a university developer club and tutor first-year IT students, experiences that have strengthened my ability "
        "to explain technical decisions, collaborate, and take ownership. I am targeting graduate and junior full-stack "
        "software roles across Ireland and am available to relocate from April 2027."
    )

    doc.add_heading("4. Recruiter and networking messages", level=1)
    doc.add_heading("Connection request", level=2)
    doc.add_paragraph(
        "Hi [Name], I expect to complete all requirements for my Software Engineering degree on 10 November 2026 and am targeting graduate full-stack "
        "roles in Ireland. My recent work includes React/Next.js projects and a live WordPress, automation, and Airtable "
        "implementation. I would value connecting and following your work in Irish technology recruitment."
    )
    doc.add_heading("Post-connection message", level=2)
    doc.add_paragraph(
        "Thanks for connecting, [Name]. I am based in Cape Town, expect to complete my degree requirements on 10 November 2026, and am available to relocate anywhere in Ireland from April 2027. I am particularly interested in graduate or junior roles using TypeScript, React, "
        "Next.js, Node.js, or cloud-backed web systems. If you recruit for this area, I would appreciate any guidance on "
        "upcoming 2027 intakes or teams able to consider employment-permit applicants."
    )
    doc.add_heading("Hiring-manager message", level=2)
    doc.add_paragraph(
        "Hi [Name], your team's work on [specific product/problem] caught my attention. I am a final-year Software "
        "Engineering student with experience delivering React/Next.js applications and real operational workflows. In a "
        "recent internship I supported a 15+ page public platform, integrated registration systems, and created technical "
        "handover resources. I am exploring graduate opportunities beginning in April 2027 and would be interested "
        "in learning what your team values most in early-career developers."
    )

    doc.add_heading("Recruiter outreach email", level=2)
    doc.add_paragraph(
        "Subject: Graduate / junior full-stack developer - available April 2027\n\n"
        "Hi [Recruiter Name], I am a Cape Town-based Software Engineering student expecting to complete all degree "
        "requirements on 10 November 2026. My experience includes React, Next.js, TypeScript, Supabase, AWS, and a "
        "live WordPress/automation/Airtable implementation for KuCoNa. I am targeting graduate or junior full-stack "
        "roles across Ireland and can relocate from April 2027. I do not currently hold Irish or EU work rights, so I "
        "would need an employer and role that can support the appropriate employment-permit process. Do you recruit "
        "for employers able to consider non-EEA early-career candidates? I would be glad to send my two-page CV and "
        "portfolio links.\n\nKind regards,\nJoshua Nehohwa"
    )

    doc.add_heading("Legitimate-recruiter checks", level=2)
    for item in [
        "Confirm the agency or its legal entity on the Workplace Relations Commission's current licensed-agency list.",
        "Start from the agency's official website; verify that the email domain and phone number match that site.",
        "Do not pay a recruiter for a job, interview, sponsorship, permit, equipment, or background check.",
        "Treat WhatsApp-only approaches, guaranteed offers, rushed bank/passport requests, and look-alike domains as warning signs.",
        "Ask directly whether the employer can consider a non-EEA applicant requiring an Irish employment permit; a recruiter cannot guarantee eligibility.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("5. Modular cover letter", level=1)
    doc.add_paragraph("Dear [Hiring Manager / Hiring Team],")
    doc.add_paragraph(
        "I am applying for the [Role] position at [Company]. I am completing a BSc in Information Technology (Software "
        "Engineering), expect to complete all degree requirements on 10 November 2026, hold an 86% academic average, and have practical experience building web "
        "applications and operational systems with [two or three technologies from the vacancy]."
    )
    doc.add_paragraph(
        "During my recent internship with Nature's Valley Trust / KuCoNa, I [choose the most relevant verified result]. "
        "This required me to [relevant problem-solving behaviour], while keeping the solution maintainable for the staff "
        "who would operate it after handover. In [project], I also [second relevant example]."
    )
    doc.add_paragraph(
        "I am drawn to [Company] because [specific product, engineering challenge, customer, or value]. The opportunity "
        "to contribute to [specific team responsibility] matches my experience in [evidence] and my goal of developing as "
        "a full-stack engineer. I am based in Cape Town and am available to relocate anywhere in Ireland from April 2027. "
        "I would require the appropriate permission to work in Ireland and would be pleased to discuss the process."
    )
    doc.add_paragraph("Thank you for considering my application. I would welcome the opportunity to discuss how my project experience, learning mindset, and ownership could contribute to [Company].")
    doc.add_paragraph("Kind regards,\nJoshua Nehohwa")

    doc.add_heading("6. Application-answer bank", level=1)
    answers = [
        ("Why are you interested in this role?", "Connect one company-specific reason, one role responsibility, and one verified example. Keep the answer between 100 and 150 words."),
        ("Tell us about a challenging project.", "Use the KuCoNa workflow or Innosimm. Explain the initial constraint, your decision, what you personally implemented, how you validated it, and the measurable result."),
        ("Do you have the right to work in Ireland?", "Answer truthfully: 'I am currently based in South Africa and do not hold Irish or EU work rights. I would require the appropriate Irish employment permission and am available to relocate from April 2027.' Do not claim an existing right to work."),
        ("Salary expectations", "State flexibility and request the approved range. Before giving a number, confirm that the basic salary meets the current permit threshold for the classified occupation."),
        ("When can you start?", "Use April 2027 as the earliest relocation/start window, subject to employment-permit processing. Degree requirements are expected complete on 10 November 2026; the ceremony is expected in May 2027."),
        ("Why Ireland?", "Focus on the engineering ecosystem, international teams, long-term professional development, and willingness to relocate—not immigration benefits alone."),
    ]
    for question, guidance in answers:
        doc.add_heading(question, level=2)
        doc.add_paragraph(guidance)

    doc.add_heading("7. Interview story bank", level=1)
    stories = [
        ("Complex integration", "KuCoNa registration workflow", "Architecture choices, privacy boundary, logging, duplicate prevention, testing and handover."),
        ("Data quality and risk", "Cleaning 1,037 records", "How scope was confirmed, integrations preserved, deletion controlled, and final state verified."),
        ("Leadership", "HackJam or Vossie DevClub", "Prioritisation, delegation, communication, resolving disagreement, and outcome."),
        ("Learning quickly", "Apps Script bilingual survey", "New tooling, constraints, iteration, verification, and user usability."),
        ("Reliability under pressure", "Snowshoe Mountain", "Safety, decisions in difficult conditions, teamwork, and recognition."),
        ("Explaining technical ideas", "Eduvos tutoring", "Adapting explanations, checking understanding, and supporting students without doing the work for them."),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    t.autofit = False
    for i, width in enumerate([1.35, 1.7, 3.55]):
        t.columns[i].width = Inches(width)
    for i, text in enumerate(["Competency", "Best example", "Points to prepare"]):
        t.rows[0].cells[i].text = text
        set_cell_shading(t.rows[0].cells[i], "007575")
        for run in t.rows[0].cells[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    set_repeat_table_header(t.rows[0])
    for values in stories:
        cells = t.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value

    doc.add_heading("8. Technical preparation checklist", level=1)
    for item in [
        "JavaScript and TypeScript: scope, closures, async/await, promises, typing, generics, error handling.",
        "React and Next.js: component design, state, hooks, rendering, forms, routing, server/client boundaries, accessibility.",
        "Backend and APIs: HTTP, REST, authentication versus authorization, validation, status codes, pagination, logging.",
        "SQL and data: joins, indexes, constraints, transactions, normalization, safe migrations, row-level security concepts.",
        "Engineering workflow: Git branching, code review, testing pyramid, debugging, CI/CD fundamentals, Docker basics.",
        "Algorithms: arrays, strings, maps/sets, stacks/queues, recursion, sorting/searching, complexity, clear tradeoff discussion.",
        "AWS: explain EC2, S3, IAM, Lambda and least privilege through realistic application examples.",
        "Portfolio: prepare a five-minute walkthrough and a deeper architecture discussion for two strongest projects.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("9. Per-application quality gate", level=1)
    for item in [
        "Vacancy saved with live URL, date captured, closing date, location, salary if stated, and full requirements.",
        "Actual duties appear consistent with an eligible software occupation; classification is never inferred from title alone.",
        "Salary and contract length are checked against current official permit rules.",
        "Every CV bullet is true, defensible, relevant, and supported by evidence.",
        "Company, role, hiring-manager name, and technology references are correct throughout.",
        "Work-authorisation answer is accurate and sponsorship is not assumed.",
        "Joshua reviews all material and explicitly approves before submission.",
        "Confirmation, status, and follow-up date are recorded immediately after submission.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("10. Current legal reference points", level=1)
    doc.add_paragraph(
        "As checked in August 2026, Ireland's Critical Skills Occupations List includes programmers and software "
        "development professionals (SOC 2136) and web design and development professionals (SOC 2137). The standard "
        "Critical Skills threshold increased to EUR40,904 on 1 March 2026. Eligibility still depends on the offered role's "
        "actual duties, salary, contract, candidate qualifications, and the rules in force when applying."
    )
    sources = [
        "https://enterprise.gov.ie/en/what-we-do/workplace-and-skills/employment-permits/employment-permit-eligibility/highly-skilled-eligible-occupations-list/",
        "https://enterprise.gov.ie/en/what-we-do/workplace-and-skills/employment-permits/permit-types/critical-skills-employment-permit/",
        "https://enterprise.gov.ie/en/news-and-events/department-news/2025/december/20251202.html",
        "https://www.irishimmigration.ie/registering-your-immigration-permission/how-to-register-your-immigration-permission-for-the-first-time/required-documents/",
    ]
    for source in sources:
        add_bullet(doc, source)

    path = OUTPUT / "Joshua_Nehohwa_Ireland_Application_Kit.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    print(build_cv())
    print(build_application_pack())
