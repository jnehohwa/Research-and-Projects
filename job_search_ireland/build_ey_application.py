from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from build_documents import (
    NAVY,
    TEAL,
    add_bullet,
    add_labelled_paragraph,
    configure_document,
)


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "outputs"
OUTPUT.mkdir(parents=True, exist_ok=True)


def add_header(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(title)
    r.font.name = "Aptos Display"
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(subtitle)
    r.font.name = "Aptos"
    r.font.size = Pt(10.2)
    r.font.bold = True
    r.font.color.rgb = TEAL


def build_ey_cv():
    doc = Document()
    configure_document(doc, compact=True)
    doc.sections[0].footer.paragraphs[0].text = ""

    add_header(doc, "JOSHUA NEHOHWA", "EY TECHNOLOGY CONSULTING GRADUATE CANDIDATE")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.add_run("Cape Town, South Africa | Open to Dublin or Cork | ")
    p.add_run("nehohwajoshua@gmail.com | +27 63 857 3965\n")
    p.add_run("linkedin.com/in/joshua-nehohwa-b4b97b229 | github.com/jnehohwa")

    doc.add_heading("PROFILE", level=1)
    doc.add_paragraph(
        "Final-year BSc Information Technology (Software Engineering) student with an 86% academic average and "
        "hands-on experience delivering web platforms, data-backed workflows, automation and stakeholder handovers. "
        "Built solutions using React, Next.js, TypeScript, PHP/MySQL, Supabase, WordPress and AWS services, including "
        "a public community platform and privacy-conscious registration workflow. Brings a consulting-oriented mix of "
        "technical problem solving, clear communication, responsible AI awareness, tutoring and team leadership. "
        "Degree requirements expected complete 10 November 2026; available for EY's September 2027 graduate intake."
    )

    doc.add_heading("TECHNOLOGY AND CONSULTING SKILLS", level=1)
    add_labelled_paragraph(doc, "Development: ", "TypeScript, JavaScript, React, Next.js, Node.js, HTML, CSS, Python, PHP, Java")
    add_labelled_paragraph(doc, "Data and platforms: ", "PostgreSQL, MySQL, Supabase, Clerk Auth, WordPress, Airtable, Jotform")
    add_labelled_paragraph(doc, "Cloud and delivery: ", "AWS EC2, S3, IAM and Lambda; Git, Linux, Docker fundamentals, Make.com, Apps Script")
    add_labelled_paragraph(doc, "Client delivery: ", "Requirements discovery, process mapping, testing, documentation, training, stakeholder communication")

    doc.add_heading("EDUCATION", level=1)
    p = doc.add_paragraph(style="CV Entry")
    p.add_run("BSc Information Technology (Software Engineering) | Eduvos, Cape Town").bold = True
    p.add_run(" | Requirements complete 10 Nov 2026")
    doc.add_paragraph(
        "Academic average: 86% | Graduation ceremony expected May 2027 | Relevant study: software architecture, "
        "data structures and algorithms, cloud technologies, network security, mobile development, mathematics and AI ethics."
    )

    doc.add_heading("RELEVANT EXPERIENCE", level=1)
    p = doc.add_paragraph(style="CV Entry")
    p.add_run("Digital Systems and Web Development Intern | Nature's Valley Trust / KuCoNa").bold = True
    p.add_run(" | Jun-Aug 2026")
    add_bullet(doc, "Translated programme and staff needs into a responsive WordPress/Elementor platform spanning three programme hubs, 15+ public pages, registration pathways and a 14-partner directory.")
    add_bullet(doc, "Designed a privacy-conscious workflow connecting WordPress, Jotform, Make.com and Airtable for registration, logging, attendance support and duplicate prevention.")
    add_bullet(doc, "Cleaned 1,037 demonstration records while preserving the production schema and integrations; documented the process and verified the final state.")
    add_bullet(doc, "Created a 47-question bilingual community survey with Apps Script and linked analytics, then produced staff training and technical handover resources.")
    add_bullet(doc, "Live platform: https://kucona.org.za")

    p = doc.add_paragraph(style="CV Entry")
    p.add_run("ITMTA1-B22 Tutor | Eduvos").bold = True
    p.add_run(" | May 2026-Present")
    add_bullet(doc, "Lead weekly two-hour Microsoft Teams tutorials, adapt explanations for different levels of understanding and support first-year IT students without replacing their own problem solving.")
    add_bullet(doc, "Coordinate exam logistics and availability with academic staff and maintain accurate tutoring-hour records.")

    doc.add_page_break()
    p = doc.add_paragraph(style="CV Entry")
    p.add_run("Freelance AI Data Specialist | Data Annotations AI").bold = True
    p.add_run(" | Remote | Jan-Nov 2024")
    add_bullet(doc, "Evaluated AI responses and transformed raw HTML and text into high-quality training annotations under complex, evolving guidelines.")
    add_bullet(doc, "Documented edge cases and actionable quality feedback while meeting independent quality and throughput expectations.")

    p = doc.add_paragraph(style="CV Entry")
    p.add_run("Seasonal Lift Attendant | Snowshoe Mountain").bold = True
    p.add_run(" | West Virginia, USA")
    add_bullet(doc, "Worked 3 Dec 2024-28 Feb 2025 and 4 Dec 2025-1 Mar 2026; contracted to return 3 Dec 2026-1 Mar 2027.")
    add_bullet(doc, "Managed safety-critical guest operations in high-pressure conditions and earned multiple top-performer recognitions for reliability and service.")

    doc.add_heading("SELECTED PROJECTS AND LEADERSHIP", level=1)
    p = doc.add_paragraph(style="CV Entry")
    p.add_run("HackJam Innovation Platform | Team Leader").bold = True
    p.add_run(" | 2025")
    add_bullet(doc, "Led a team building a React, Next.js and TypeScript platform for idea submission, voting, mentor feedback and gamification; placed 4th university-wide.")
    add_bullet(doc, "Coordinated user research, prototyping, technical priorities and the final stakeholder pitch. Repository: https://github.com/jnehohwa/HackJam-Shark-Tank")

    p = doc.add_paragraph(style="CV Entry")
    p.add_run("KasiSwap Marketplace | Developer").bold = True
    p.add_run(" | 2026")
    add_bullet(doc, "Built React/TypeScript and PHP/MySQL flows for listings, authentication-aware actions, orders, messaging, disputes, reviews and administration.")
    add_bullet(doc, "Produced deployment notes, code evidence and maintainable technical documentation. Demo: https://kasiswap.free.nf/?i=1")

    p = doc.add_paragraph(style="CV Entry")
    p.add_run("Innosimm ERP System | Developer").bold = True
    p.add_run(" | 2025-Present")
    add_bullet(doc, "Designing a modular Next.js, Supabase and Clerk system to replace fragmented inventory, sales, purchasing and finance processes; currently in development.")

    p = doc.add_paragraph(style="CV Entry")
    p.add_run("Chairperson | Vossie DevClub").bold = True
    p.add_run(" | 2025-Present")
    add_bullet(doc, "Lead workshops, mentoring, hackathons and programming outreach; communicate technical ideas to audiences with varied experience.")

    doc.add_heading("CERTIFICATIONS AND ACHIEVEMENTS", level=1)
    add_bullet(doc, "AWS Certified Solutions Architect - Associate | https://www.credly.com/badges/43660ab1-f99a-4bbe-ab14-261dfb57cfcd")
    add_bullet(doc, "AWS Certified Cloud Practitioner | https://www.credly.com/badges/61737c5d-678d-4633-80d4-7ea1c5d24b80")
    add_bullet(doc, "Golden Key International Honour Society | Top 15% of university cohort, 2025")
    add_bullet(doc, "IELTS Academic: 8.5 overall / CEFR C2")

    path = OUTPUT / "Joshua_Nehohwa_EY_Technology_Consulting_CV.docx"
    doc.save(path)
    return path


def add_answer(doc, question, answer):
    doc.add_heading(question, level=2)
    doc.add_paragraph(answer)


def build_ey_answers():
    doc = Document()
    configure_document(doc, compact=False)
    add_header(doc, "EY TECHNOLOGY CONSULTING APPLICATION PACK", "Joshua Nehohwa | Graduate Programme 2027")

    doc.add_heading("APPLICATION TARGET", level=1)
    add_labelled_paragraph(doc, "Programme: ", "EY Ireland - Technology Consulting Graduate Programme 2027")
    add_labelled_paragraph(doc, "Preferred location: ", "Dublin")
    add_labelled_paragraph(doc, "Preferred client group: ", "All Industry Clients")
    add_labelled_paragraph(doc, "Areas of strongest fit: ", "Systems Engineering; Microsoft Cloud & BizApps; Digital Assurance and Testing; AI & Data")
    add_labelled_paragraph(doc, "Availability: ", "September 2027 intake")

    doc.add_heading("TAILORED MOTIVATION ANSWERS", level=1)
    add_answer(doc, "Why do you want to join EY Technology Consulting?", (
        "I am interested in EY Technology Consulting because it combines technology delivery with the responsibility of understanding a client's real operating problem. "
        "That is the kind of work I enjoyed during my KuCoNa internship, where I translated staff and programme needs into a public website, registration workflows, data controls and practical handover resources. "
        "EY's work across technology strategy, full-stack implementation, cloud, data, testing and emerging technologies would let me build on my development experience while learning how complex organisations make and deliver technology decisions. "
        "I am particularly attracted to the graduate induction, continuous learning and access to multidisciplinary teams. I would bring curiosity, disciplined problem solving, clear communication and a willingness to learn from feedback while contributing practical experience with React, Next.js, TypeScript, automation, data platforms and AWS."
    ))
    add_answer(doc, "Why Technology Consulting rather than a purely software-development role?", (
        "I enjoy writing software, but I am most motivated when the technical work solves a visible user or organisational problem. At KuCoNa, the value was not simply launching pages or connecting tools; it was understanding registration, privacy, duplicate prevention, analytics and what staff needed to operate the system after handover. "
        "Technology Consulting offers that broader problem-solving environment. It requires listening carefully, structuring an unclear problem, choosing an appropriate solution, communicating trade-offs and helping users adopt the result. My development background gives me enough technical depth to participate credibly in implementation, while tutoring, student leadership and international customer-facing work have strengthened how I explain ideas and work with different people."
    ))
    add_answer(doc, "What interests you about AI and emerging technology?", (
        "I am interested in AI as a practical tool that must be applied responsibly and verified carefully. In freelance AI-data work, I evaluated model responses, transformed raw material into structured annotations and documented edge cases under evolving guidelines. That experience showed me that useful AI depends on data quality, clear evaluation criteria and human judgement. "
        "My degree also covers AI ethics, and I use AI as a collaborative tool while retaining responsibility for requirements, technical decisions, testing and final outputs. In consulting, I would look for problems where AI can improve analysis, automation or user experience, while considering privacy, accuracy, security and whether a simpler solution would be more appropriate."
    ))
    add_answer(doc, "Tell us about a challenging project.", (
        "During my internship with Nature's Valley Trust and KuCoNa, I helped deliver a public platform and registration workflow across WordPress, Jotform, Make.com and Airtable. The challenge was that the solution needed to support several programme areas, protect personal information, prevent duplicates and remain understandable to staff after the internship. "
        "I mapped the workflow, separated public collection from operational records, added integration logging and duplicate-prevention logic, and documented the final process. I also cleaned 1,037 demonstration records without damaging the production schema or integrations. The result was a live 15+ page platform, working registration pathways and handover resources. The experience taught me to treat maintainability, controls and user adoption as part of the technical solution rather than as afterthoughts."
    ))
    add_answer(doc, "Describe a time you led a team.", (
        "I led a student team during HackJam to build a platform for idea submission, voting, mentor feedback and gamification. The team had to balance user expectations, limited development time and the need to present a coherent working concept. I helped organise user research, break the product into priorities, coordinate responsibilities and keep the technical work aligned with the final pitch. "
        "When ideas competed for attention, I brought the discussion back to user value and what we could demonstrate reliably. We delivered the React, Next.js and TypeScript solution and placed fourth university-wide. I learned that leadership is less about having every answer and more about creating clarity, listening, making timely decisions and helping each person contribute effectively."
    ))
    add_answer(doc, "How do you work inclusively with people who have different levels of technical knowledge?", (
        "As an Eduvos tutor and Vossie DevClub chairperson, I regularly explain technical ideas to students with different backgrounds and confidence levels. I start by asking what they already understand, use a concrete example, and then check their reasoning rather than simply giving them an answer. "
        "During the KuCoNa handover, I also had to present operational workflows in a way that staff could maintain without needing developer-level knowledge. These experiences taught me to avoid unnecessary jargon, invite questions and treat different perspectives as useful information about whether a solution is genuinely understandable."
    ))

    doc.add_heading("SCREENING FACTS", level=1)
    add_answer(doc, "Do you have the right to work in Ireland?", (
        "I am a South African citizen and do not currently hold Irish or EU work rights. I would require the appropriate Irish employment permit and immigration permission. EY's published graduate FAQ identifies Consulting and Technology as a programme that can support eligible non-EEA applicants. I understand that I must satisfy the applicable criteria and have valid permission in place before starting."
    ))
    add_answer(doc, "When can you start?", (
        "I will complete all degree modules and credits on 10 November 2026. My formal graduation ceremony is expected in May 2027. I am available for EY's main September 2027 graduate intake and can relocate earlier if required for the employment-permit and onboarding process."
    ))
    add_answer(doc, "Academic result", (
        "Current academic average: 86%. Degree: BSc Information Technology (Software Engineering), Eduvos, Cape Town. Degree requirements are expected to be complete on 10 November 2026."
    ))
    add_answer(doc, "Location and client preference", (
        "Dublin and All Industry Clients are my recommended first preferences because they provide the broadest exposure across technology, government, infrastructure, energy, consumer and private-sector transformation. Cork remains a suitable alternative."
    ))

    doc.add_heading("OPTIONAL COVER LETTER", level=1)
    doc.add_paragraph("Dear EY Graduate Recruitment Team,")
    doc.add_paragraph(
        "I am applying for the EY Ireland Technology Consulting Graduate Programme 2027. I am completing a BSc in Information Technology (Software Engineering) at Eduvos, where I hold an 86% academic average, and I expect to complete all degree requirements on 10 November 2026."
    )
    doc.add_paragraph(
        "My strongest experience sits where technical delivery meets a real operational need. During my internship with Nature's Valley Trust and KuCoNa, I helped build a 15+ page public platform and a registration workflow connecting WordPress, Jotform, Make.com and Airtable. I considered privacy, duplicate prevention, data quality and staff handover alongside implementation, and safely removed 1,037 demonstration records while preserving the production schema and integrations."
    )
    doc.add_paragraph(
        "EY appeals to me because Technology Consulting spans strategic advice, full-stack implementation, cloud, data, systems engineering and emerging technology. I would bring practical development experience, AWS foundations, responsible AI awareness and the communication skills developed through tutoring, student leadership and international work. I am especially interested in helping multidisciplinary teams turn complex client needs into maintainable, well-tested solutions."
    )
    doc.add_paragraph(
        "I am a South African citizen and would require the appropriate Irish employment permission. I am available for the September 2027 intake and am prepared to relocate to Ireland. Thank you for considering my application."
    )
    doc.add_paragraph("Kind regards,\nJoshua Nehohwa")

    doc.add_heading("FINAL REVIEW BEFORE SUBMISSION", level=1)
    for item in [
        "Confirm that the application form accepts an overseas degree and enter the qualification exactly as shown on the transcript.",
        "Use Technology Consulting as the single EY programme application; EY does not allow simultaneous applications to multiple student programmes.",
        "Select Dublin and All Industry Clients unless Joshua prefers Cork or Global Financial Services after reviewing the form.",
        "Upload the EY-tailored two-page CV, not the general Ireland CV.",
        "Answer work-authorisation questions exactly as written; do not state that sponsorship or a permit is already guaranteed.",
        "Review any online-assessment declarations and disability-adjustment question personally.",
        "Do not submit until Joshua has reviewed the complete form and explicitly approved submission.",
    ]:
        add_bullet(doc, item)

    path = OUTPUT / "Joshua_Nehohwa_EY_Application_Answers.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    print(build_ey_cv())
    print(build_ey_answers())
